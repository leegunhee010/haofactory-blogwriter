# -*- coding: utf-8 -*-
"""
하오팩토리 블로그 작성기 — Claude Code(각자 Max/Pro 구독)로 글 생성
=====================================================================
- 탭 여러 개 = 한 번에 여러 글
- 키워드 직접 입력 또는 통검 글감추천에서 자동삽입
- 사진 폴더 지정 → 그 안 사진을 (사진N) 자리에 자동 배치
- 생성 → 미리보기 → 워드 저장(사진 박힌 채) → 네이버 복붙
엔진: claude.exe -p (각자 자기 구독으로 인증, API 종량제 아님)
실행: 작성기_실행.bat → http://127.0.0.1:5002
"""
import os, sys, json, glob, shutil, subprocess, threading, time, webbrowser, re, datetime, random
from flask import Flask, request, jsonify, send_file

FROZEN = getattr(sys, "frozen", False)
HERE = os.path.dirname(sys.executable) if FROZEN else os.path.dirname(os.path.abspath(__file__))
OUT_DIR = os.path.join(HERE, "출력")
os.makedirs(OUT_DIR, exist_ok=True)
CFG_FILE = os.path.join(HERE, "작성기_설정.json")
TEMPLATE = os.path.join(HERE, "템플릿", "카드뉴스_템플릿.pptx")
sys.path.insert(0, HERE)
import cardnews_pil  # noqa  (순수 Pillow 엔진 — PowerPoint 불필요)
import brands  # noqa  (브랜드 설정 관리)

CARD_JOBS = {}
CARD_LOCK = threading.Lock()


def _cardnews_job(jid, photo_paths, cards, keyword, assets_dir, subtitle="", bodies=None, title=""):
    try:
        safe = re.sub(r'[\\/:*?"<>|]', "_", keyword)
        stamp = datetime.datetime.now().strftime("%H%M%S%f")[:9]
        png_dir = os.path.join(OUT_DIR, f"{safe}_카드뉴스_{stamp}")
        res = cardnews_pil.make_cards(photo_paths, cards[:7], png_dir, assets_dir,
                                      subtitle=subtitle, bodies=bodies, title=title)   # 순수 PIL, 즉시
        srcs = [c["src"] for c in res["cards"]]
        CARD_JOBS[jid] = {"status": "done", "pngs": res["pngs"], "dir": png_dir, "srcs": srcs,
                          "cards": res["cards"]}
    except Exception as e:
        CARD_JOBS[jid] = {"status": "error", "msg": str(e)[:200]}

# 통검 데이터(글감추천용) 기본 경로 — 하오팩토리 통검체크 dist
DEFAULT_TONGGEOM = r"C:/Users/이건희/trevity-노출체크/dist/하오팩토리 통검체크"

app = Flask(__name__)


def jload(p, d):
    try:
        return json.load(open(p, encoding="utf-8"))
    except Exception:
        return d


def load_cfg():
    c = jload(CFG_FILE, {})
    c.setdefault("tonggeom_dir", DEFAULT_TONGGEOM)
    c.setdefault("brand_name", "하오팩토리")
    c.setdefault("model", "opus")  # 품질 우선(기본). sonnet은 더 빠름
    return c


# ── claude.exe 자동 탐지 ────────────────────────────────
def find_claude():
    p = shutil.which("claude")
    if p:
        return p
    pats = [
        os.path.expanduser(r"~/AppData/Local/Packages/Claude_*/LocalCache/Roaming/Claude/claude-code/*/claude.exe"),
        os.path.expanduser(r"~/.local/bin/claude.exe"),
        os.path.expanduser(r"~/AppData/Roaming/npm/claude.cmd"),
    ]
    hits = []
    for g in pats:
        hits += glob.glob(g)
    # 버전 숫자 큰 것(최신) 우선
    hits.sort()
    return hits[-1] if hits else None


CLAUDE = find_claude()


# ── 멀티 Claude 계정 (CLAUDE_CONFIG_DIR로 계정별 분리) ───
# 각 계정은 accounts/<id>/ 폴더 = 그 계정의 CLAUDE_CONFIG_DIR(자격증명 격리).
# "default" 계정은 dir 비움 → 기존 ~/.claude 그대로 사용(앱 최초 동작 보존).
ACCOUNTS_DIR = os.path.join(HERE, "accounts")
ACCOUNTS_FILE = os.path.join(HERE, "accounts.json")
# 사용 한도/요금 한도 도달 메시지 패턴(자동 다른 계정 전환 트리거)
_LIMIT_PAT = re.compile(
    r"(spend limit|usage limit|rate limit|monthly .{0,20}limit|hit your .{0,20}limit|"
    r"limit.{0,40}claude\.ai|reached your|out of .{0,15}usage|사용 한도|한도에 도달|한도를 초과)", re.I)


def load_accounts():
    d = jload(ACCOUNTS_FILE, {})
    accts = d.get("accounts") or []
    if not any(a.get("id") == "default" for a in accts):
        accts.insert(0, {"id": "default", "label": "기본 계정", "dir": ""})
    active = d.get("active") or "default"
    if not any(a.get("id") == active for a in accts):
        active = "default"
    return {"accounts": accts, "active": active}


def save_accounts(d):
    try:
        json.dump(d, open(ACCOUNTS_FILE, "w", encoding="utf-8"), ensure_ascii=False, indent=1)
    except Exception:
        pass


def account_env(acc_id=None):
    """해당 계정의 CLAUDE_CONFIG_DIR 적용된 env. dir 비면 기본(~/.claude)."""
    d = load_accounts()
    if acc_id is None:
        acc_id = d["active"]
    acc = next((a for a in d["accounts"] if a.get("id") == acc_id), None)
    env = os.environ.copy()
    if acc and acc.get("dir"):
        os.makedirs(acc["dir"], exist_ok=True)
        env["CLAUDE_CONFIG_DIR"] = acc["dir"]
    else:
        env.pop("CLAUDE_CONFIG_DIR", None)
    return env


def auth_status_of(acc_id):
    """계정별 로그인 상태(dict: loggedIn/email/subscriptionType ...)."""
    if not CLAUDE:
        return {}
    try:
        r = subprocess.run([CLAUDE, "auth", "status"], capture_output=True, text=True,
                           encoding="utf-8", errors="ignore", timeout=25, env=account_env(acc_id))
        return json.loads(r.stdout or "{}")
    except Exception:
        return {}


# ── 하오팩토리 문체 프롬프트 ─────────────────────────────
STYLE = """너는 조형물 제작 회사 '하오팩토리(HAOFACTORY)'의 네이버 블로그 글을 쓰는 전문 카피라이터다.
[문체 규칙]
- 차분하고 현장감 있는 실무자 톤. 과장·AI티('~에 대해 알아보겠습니다' 등)·이모지 금지.
- 한 문장(또는 의미단위)마다 줄바꿈(네이버 모바일 가독성). 종결은 ~합니다/~됩니다 위주.
- 도입: "안녕하세요." + 한 줄 소개 + "하오팩토리입니다." 로 시작.
- 넘버링 소제목 **정확히 6개**(개요/요구사항/설계/공정/설치/결과 흐름). 소재 비교(금속·FRP·EPS) 자연스럽게.
- 마무리: "~ 고민되는 경우라면 HAOFACTORY로 편하게 문의 주셔도 좋습니다."
[SEO 규칙]
- 공백 제외 1,500자 이상.
- 제목에 메인키워드 포함. 본문에 메인키워드 5~8회 자연스럽게(억지 반복 금지).
[출력 형식 — 정확히 이 형식만, 다른 설명/머리말 금지]
제목: (메인키워드 포함한 제목)
(사진1: 파일명)
본문 도입...
**1. 소제목**
(사진2: 파일명)
본문...
(... 정확히 6개 소제목, 각 소제목 아래 사진 1장씩 ...)
★(사진)은 **정확히 7개**: 도입 1장(사진1) + 소제목 6개 각 1장(사진2~사진7). 카드뉴스 7장과 1:1로 맞춘다. 더도 덜도 말 것.
규칙: 사진은 아래 '사용 사진' 목록의 파일명만 (사진N: 정확한파일명) 형태로 배치. 목록에 없는 파일명 만들지 말 것.
★중요: **각 (사진N) 슬롯에는 서로 다른 사진**을, 그 단락 내용(개요/설계/공정/설치/완성 등)에 **가장 어울리는 사진**으로 배치하라. 같은 사진을 두 번 쓰지 말 것. 파일명에 단계 힌트(대표/공정/설치/완성/도장 등)가 있으면 그 단락에 맞춰라.

[카드뉴스 헤드라인]
원고를 다 쓴 뒤, 맨 끝에 아래 형식으로 카드 7장 헤드라인을 출력하라(짧고 강하게, 각 12자 내외, 본문 흐름 요약):
카드뉴스:
1. (메인키워드 그대로 — 표지)
2. (포토존/공간 후킹)
3. (기획/설계 이유)
4. (안전성/구조)
5. (소재/입체감)
6. (설치/현장)
7. (전문회사 선택 마무리)"""


# 매 생성마다 '표현·각도'만 살짝 바꾸는 변주 풀(구조·톤은 건드리지 않음)
_VAR_HOOK = ["고객이 자주 하는 고민·질문을 짚으며", "흔한 오해 하나를 바로잡으며",
             "실제 현장 장면을 그리듯", "핵심 결론을 먼저 던진 뒤 풀어가며",
             "예전 방식과 지금을 비교하며", "비슷한 고민을 했던 사례를 떠올리게 하며"]
_VAR_SUB = ["질문형 소제목 위주로", "단정·선언형 소제목 위주로",
            "호기심을 자극하는 소제목으로", "핵심 키워드 + 짧은 구절형 소제목으로"]
_VAR_EMPH = ["왜 필요한지(필요성)", "흔한 실수·실패 사례", "선택 기준과 체크포인트",
             "작업 과정과 디테일", "결과와 효과", "자주 묻는 질문 관점", "비용·기간의 현실"]
_VAR_RHYTHM = ["사례 중심으로", "실용 정보·팁 중심으로", "공감·스토리 중심으로"]


def _variation_block():
    emph = random.sample(_VAR_EMPH, 2)
    return ("[이번 글 변주 — 아래는 '표현·각도'만 새롭게 하기 위한 지시다. "
            "블로그 구조(도입 인사 → 소제목 정확히 6개 → 마무리 CTA), 키워드 SEO 규칙, "
            "브랜드 톤·문체·정체성은 절대 바꾸지 말 것. 틀과 기본기는 그대로, 표현만 매번 다르게.]\n"
            f"- 도입은 {random.choice(_VAR_HOOK)} 자연스럽게 시작한다.\n"
            f"- 소제목은 {random.choice(_VAR_SUB)} 구성한다(매번 똑같은 문형 반복은 피한다).\n"
            f"- 이번 글은 특히 '{emph[0]}'·'{emph[1]}' 관점을 조금 더 살린다.\n"
            f"- 전체 전개는 {random.choice(_VAR_RHYTHM)} 풀어간다.\n"
            "- 같은 키워드라도 지난 글과 같은 첫 문장·같은 예시·같은 비유의 반복을 피하고 새 각도로 쓴다.")


def build_prompt(keyword, photo_files, project_hint="", brand=None):
    photos = "\n".join(f"- {f}" for f in photo_files) if photo_files else "(없음 — (사진N) 자리표시만, 파일명은 비워둠)"
    hint = f"\n[실제 프로젝트 정보] {project_hint}" if project_hint else ""
    style = brands.build_style(brand) if brand else STYLE
    return (style + "\n\n" + _variation_block() +
            f"\n\n[메인키워드] {keyword}{hint}\n\n[사용 사진] (순서대로 적절한 슬롯에 배치)\n{photos}\n\n위 규칙대로 지금 작성하라.")


def run_claude(prompt, model="", acc_id=None, _tried=None):
    """활성 계정으로 claude -p 실행. 사용 한도에 걸리면 로그인된 다른 계정으로 자동 전환."""
    if not CLAUDE:
        return None, "claude 실행파일을 찾지 못했습니다. Claude 데스크톱 앱 또는 CLI를 설치하세요."
    d = load_accounts()
    if acc_id is None:
        acc_id = d["active"]
    _tried = _tried if _tried is not None else []
    _tried.append(acc_id)
    cmd = [CLAUDE, "-p"]
    if model:
        cmd += ["--model", model]
    try:
        r = subprocess.run(cmd, input=prompt, capture_output=True, text=True,
                           encoding="utf-8", errors="ignore", timeout=240, env=account_env(acc_id))
        out = (r.stdout or "").strip()
        combined = out + "\n" + (r.stderr or "")
        if "Not logged in" in combined or "Please run /login" in combined:
            return None, "이 계정은 Claude 로그인이 필요합니다. 상단 👤계정에서 로그인하세요."
        if _LIMIT_PAT.search(combined):
            # 한도 도달 → 로그인된 다른 계정으로 자동 전환
            for a in d["accounts"]:
                if a["id"] in _tried:
                    continue
                if auth_status_of(a["id"]).get("loggedIn"):
                    o2, e2 = run_claude(prompt, model, a["id"], _tried)
                    if o2 is not None and a["id"] != d["active"]:  # 작동한 계정으로 활성 고정
                        d["active"] = a["id"]; save_accounts(d)
                        _AUTH_CACHE["t"] = 0
                    return o2, e2
            return None, ("모든 계정이 사용 한도에 도달했습니다. "
                          "상단 👤계정에서 다른 Claude 계정을 추가하거나, claude.ai에서 한도를 올려주세요.")
        if not out:
            return None, "빈 응답입니다. " + (r.stderr or "")[:200]
        return out, None
    except subprocess.TimeoutExpired:
        return None, "생성 시간이 초과됐습니다(240초). 다시 시도하세요."
    except Exception as e:
        return None, f"실행 오류: {e}"


# ── 원고 파싱 (제목 + 본문블록 + 사진슬롯) ───────────────
def parse_manuscript(text, photo_files):
    lines = text.split("\n")
    title = ""
    blocks = []  # {"type":"text"|"photo", ...}
    nameset = {f.lower(): f for f in photo_files}
    buf = []

    def flush():
        if buf:
            t = "\n".join(buf).strip("\n")
            if t.strip():
                blocks.append({"type": "text", "text": t})
        buf.clear()

    for ln in lines:
        s = ln.strip()
        m_title = re.match(r"^제목\s*[:：]\s*(.+)$", s)
        m_photo = re.match(r"^\(사진[^:：]*[:：]\s*(.*?)\)\s*$", s)
        if m_title and not title:
            title = m_title.group(1).strip()
            continue
        if m_photo:
            flush()
            fn = m_photo.group(1).strip()
            real = nameset.get(fn.lower())
            if not real:  # 부분매칭 시도
                for k, v in nameset.items():
                    if fn and (fn.lower() in k or k in fn.lower()):
                        real = v; break
            blocks.append({"type": "photo", "file": real or "", "raw": fn})
            continue
        buf.append(ln)
    flush()
    # 제목 못 찾으면 첫 텍스트 줄
    if not title:
        for b in blocks:
            if b["type"] == "text":
                title = b["text"].split("\n")[0][:60]; break
    # 글자수
    body = "".join(b["text"] for b in blocks if b["type"] == "text")
    nospace = re.sub(r"\s", "", body)
    return {"title": title, "blocks": blocks, "char_count": len(nospace)}


# ── 라우트 ───────────────────────────────────────────────
_AUTH_CACHE = {"t": 0, "v": False, "id": None}


def claude_logged_in():
    if not CLAUDE:
        return False
    active = load_accounts()["active"]
    if (_AUTH_CACHE["v"] and _AUTH_CACHE["id"] == active
            and time.time() - _AUTH_CACHE["t"] < 180):  # 같은 계정 로그인됨이면 3분 캐시
        return True
    v = bool(auth_status_of(active).get("loggedIn"))
    _AUTH_CACHE.update(t=time.time(), v=v, id=active)
    return v


@app.route("/api/status")
def api_status():
    return jsonify(claude=bool(CLAUDE), claude_path=CLAUDE or "", cfg=load_cfg())


@app.route("/api/auth-status")
def api_auth_status():
    return jsonify(logged_in=claude_logged_in())


@app.route("/api/login", methods=["POST"])
def api_login():
    """활성(또는 지정) 계정으로 로그인 콘솔 띄우기."""
    if not CLAUDE:
        return jsonify(ok=False, msg="claude 실행파일을 찾지 못했습니다.")
    aid = (request.get_json(silent=True) or {}).get("id")
    try:
        flags = 0x10 if os.name == "nt" else 0   # CREATE_NEW_CONSOLE
        subprocess.Popen([CLAUDE, "auth", "login"], creationflags=flags, env=account_env(aid))
        _AUTH_CACHE["t"] = 0
        return jsonify(ok=True)
    except Exception as e:
        return jsonify(ok=False, msg=str(e))


# ── 멀티 계정 라우트 ─────────────────────────────────────
@app.route("/api/accounts")
def api_accounts():
    d = load_accounts()
    out = []
    for a in d["accounts"]:
        st = auth_status_of(a["id"]) if CLAUDE else {}
        out.append({"id": a["id"], "label": a.get("label", a["id"]),
                    "logged_in": bool(st.get("loggedIn")),
                    "email": st.get("email", ""),
                    "plan": st.get("subscriptionType", "")})
    return jsonify(accounts=out, active=d["active"])


@app.route("/api/account-add", methods=["POST"])
def api_account_add():
    """새 계정 슬롯 생성 + 그 계정 전용 로그인 콘솔 띄우기."""
    if not CLAUDE:
        return jsonify(ok=False, msg="claude 실행파일을 찾지 못했습니다.")
    b = request.get_json(force=True) or {}
    label = (b.get("label") or "").strip() or "추가 계정"
    d = load_accounts()
    existing = {a["id"] for a in d["accounts"]}
    n = 2
    while f"acc{n}" in existing:
        n += 1
    aid = f"acc{n}"
    adir = os.path.join(ACCOUNTS_DIR, aid)
    os.makedirs(adir, exist_ok=True)
    d["accounts"].append({"id": aid, "label": label, "dir": adir})
    # active는 바꾸지 않음 — 로그인 완료가 확인되면 프런트가 자동 전환(중도 취소 시 앱 안 멈춤)
    save_accounts(d)
    _AUTH_CACHE["t"] = 0
    try:
        flags = 0x10 if os.name == "nt" else 0
        subprocess.Popen([CLAUDE, "auth", "login"], creationflags=flags, env=account_env(aid))
    except Exception as e:
        return jsonify(ok=False, msg=str(e))
    return jsonify(ok=True, id=aid)


@app.route("/api/account-switch", methods=["POST"])
def api_account_switch():
    b = request.get_json(force=True) or {}
    aid = b.get("id")
    d = load_accounts()
    if not any(a["id"] == aid for a in d["accounts"]):
        return jsonify(ok=False, msg="없는 계정입니다.")
    d["active"] = aid
    save_accounts(d)
    _AUTH_CACHE["t"] = 0
    return jsonify(ok=True)


@app.route("/api/account-remove", methods=["POST"])
def api_account_remove():
    b = request.get_json(force=True) or {}
    aid = b.get("id")
    if aid == "default":
        return jsonify(ok=False, msg="기본 계정은 삭제할 수 없습니다.")
    d = load_accounts()
    acc = next((a for a in d["accounts"] if a["id"] == aid), None)
    d["accounts"] = [a for a in d["accounts"] if a["id"] != aid]
    if d["active"] == aid:
        d["active"] = "default"
    save_accounts(d)
    _AUTH_CACHE["t"] = 0
    try:
        if acc and acc.get("dir") and os.path.isdir(acc["dir"]):
            shutil.rmtree(acc["dir"], ignore_errors=True)
    except Exception:
        pass
    return jsonify(ok=True)


@app.route("/api/browse")
def api_browse():
    """로컬 폴더 탐색 — 업로드 없이 서버가 직접 읽음."""
    path = request.args.get("path", "")
    exts = (".jpg", ".jpeg", ".png", ".webp", ".gif")
    if not path:  # 드라이브 + 바로가기
        import string
        drives = [d + ":\\" for d in string.ascii_uppercase if os.path.exists(d + ":\\")]
        home = os.path.expanduser("~")
        places = []
        for label, sub in [("🖥 바탕화면", "Desktop"), ("⬇ 다운로드", "Downloads"),
                            ("📄 문서", "Documents"), ("🖼 사진", "Pictures")]:
            p = os.path.join(home, sub)
            if os.path.isdir(p):
                places.append({"name": label, "path": p})
        places.append({"name": "🏠 내 폴더", "path": home})
        return jsonify(path="", parent="", folders=drives, images=0, drive=True, places=places)
    folders, images = [], 0
    try:
        for name in sorted(os.listdir(path), key=lambda s: s.lower()):
            full = os.path.join(path, name)
            try:
                if os.path.isdir(full):
                    folders.append(name)
                elif name.lower().endswith(exts):
                    images += 1
            except Exception:
                pass
    except Exception as e:
        return jsonify(error=str(e), path=path)
    parent = os.path.dirname(path.rstrip("\\/")) if path.rstrip("\\/") != path.rstrip("\\/")[:2] else ""
    return jsonify(path=path, parent=parent, folders=folders, images=images, drive=False)


def list_images(folder, cap=20000):
    """폴더(+하위 폴더 전부)의 이미지를 폴더기준 상대경로로. 상위폴더 지정해도 하위 사진 전부 다 찾음."""
    exts = (".jpg", ".jpeg", ".png", ".webp", ".gif")
    files = []
    if folder and os.path.isdir(folder):
        for root, dirs, fs in os.walk(folder):
            for fn in fs:
                if fn.lower().endswith(exts):
                    files.append(os.path.relpath(os.path.join(root, fn), folder).replace("\\", "/"))
            if len(files) > cap * 3:
                break
    files.sort()
    return files[:cap]


GENERIC_KW = {"조형물", "제작", "제작업체", "업체", "회사", "전문", "디자인", "사례", "포토존"}


def keyword_terms(kw, type_words=None):
    """키워드에서 소재/유형 핵심어 추출. 유형단어가 키워드에 박혀 있으면 그걸로.
    예) 스티로폼스카시 → [스티로폼, 스카시], FRP조형물 → [FRP]."""
    tw = type_words if type_words is not None else _TYPE_WORDS
    found = [t for t in tw if t.lower() in kw.lower()]
    if found:
        return found
    terms = []
    for t in re.split(r"\s+", kw.strip()):
        t2 = t.replace("조형물", "").replace("제작", "").strip()
        if t2 and t2 not in GENERIC_KW and len(t2) >= 2:
            terms.append(t2)
    return terms


def filter_by_keyword(files, kw, type_words=None):
    """키워드 핵심어가 든 사진 매칭. **전용 사진 7장+면 그것만**(예: 스카시),
    부족하면(예: FRP=5장뿐) 전체 라이브러리에서 다양하게 — 같은 사진 반복 방지."""
    terms = keyword_terms(kw, type_words)
    if not terms:
        return files
    matched = [f for f in files if any(tm.lower() in f.lower() for tm in terms)]
    if len(matched) >= 7:
        return matched                      # 전용 사진 충분(카드 7장) → 키워드 사진만
    return files if len(files) > len(matched) else matched   # 전용 적으면(예: FRP 5장) 전체에서 다양하게


def _scene_key(path):
    """같은 장면(연속컷) 묶음 키 — 파일명 끝 번호 제거. [날짜]_[프로젝트]_[단계]_[번호] → 번호빼고 그룹."""
    base = os.path.splitext(os.path.basename(path))[0].lower()
    base = re.sub(r"[ _\-]*\(?\d+\)?$", "", base)   # 끝 번호/(번호) 제거
    return base or os.path.dirname(path).lower()


def _project_key(path):
    """프로젝트(최상위 폴더) 키 — 카드를 서로 다른 프로젝트에서 뽑기 위해."""
    norm = path.replace("\\", "/").strip("/")
    parts = norm.split("/")
    if len(parts) > 1:
        return parts[0].lower()
    return (os.path.basename(os.path.dirname(path)) or path).lower()


def pick_diverse(all_files, prefer, n, rnd=None):
    """카드용 사진 n장. **Claude가 본문 단락에 매칭한 사진(prefer)을 순서대로 우선** → 카드 헤드라인과 사진이 같은 섹션.
    모자라면 다양한 사진으로 채움(프로젝트 라운드로빈 + 같은 장면 회피). 변주는 생성 시 사진 셔플로 유지됨."""
    rnd = rnd or random.Random()
    allset = set(all_files)
    chosen = []
    # 1) Claude가 각 단락에 고른 사진을 본문 순서 그대로(★장면 중복이어도 빼지 않음 — 그래야 카드-단락 1:1)
    for f in (prefer or []):
        if f in allset and f not in chosen:
            chosen.append(f)
            if len(chosen) >= n:
                return chosen[:n]
    # 2) 부족분만 다양한 사진으로 채움(랜덤 + 프로젝트별 라운드로빈 + 같은 장면 회피)
    seen = set(_scene_key(f) for f in chosen)   # dedup은 채움분에만 적용
    pool = list(dict.fromkeys(all_files))
    rnd.shuffle(pool)
    groups = {}                               # 프로젝트(최상위 폴더)별 묶기
    for f in pool:
        groups.setdefault(_project_key(f), []).append(f)
    order = list(groups.keys()); rnd.shuffle(order)
    gi = {k: 0 for k in groups}
    # 라운드로빈: 프로젝트를 돌아가며 한 장씩(서로 다른 주제), 같은 장면은 건너뜀
    while len(chosen) < n and any(gi[k] < len(groups[k]) for k in order):
        for k in order:
            if len(chosen) >= n:
                break
            while gi[k] < len(groups[k]):
                f = groups[k][gi[k]]; gi[k] += 1
                if f in chosen or _scene_key(f) in seen:
                    continue
                seen.add(_scene_key(f)); chosen.append(f); break
    if len(chosen) < n:                       # 부족하면 남은 것으로 채움(장면 중복 허용)
        for f in pool:
            if f not in chosen:
                chosen.append(f)
                if len(chosen) >= n:
                    break
    return chosen[:n]


@app.route("/api/photos")
def api_photos():
    folder = request.args.get("folder", "")
    files = list_images(folder)
    return jsonify(folder=folder, files=files)


# ── 사진 정리 (폴더명·파일명·비전으로 분류 → 원본 이름변경) ──────────
ORG = {"running": False, "total": 0, "done": 0, "renamed": 0, "log": [], "finished": False}
_STAGE_HINTS = [("대표", "대표"), ("콘셉트", "대표"), ("조각", "공정"), ("몰드", "공정"),
                ("적층", "공정"), ("샌딩", "공정"), ("퍼티", "공정"), ("공정", "공정"),
                ("도장", "도장"), ("채색", "도장"), ("설치", "설치"), ("현장", "설치"),
                ("완성", "완성"), ("디테일", "완성")]
_TYPE_WORDS = ["스카시", "금속", "캐릭터", "마스코트", "FRP", "EPS", "스티로폼", "조각",
               "피규어", "동상", "글씨", "간판", "팝업", "포토존"]


def _stage_from_name(name, stages=None):
    for k, v in (stages if stages is not None else _STAGE_HINTS):
        if k in name:
            return v
    return None


def _type_from_path(rel, type_words=None):
    low = rel.lower()
    for t in (type_words if type_words is not None else _TYPE_WORDS):
        if t.lower() in low:
            return t
    return ""


def _campaign_from_path(rel):
    """부모 폴더명 = 캠페인/프로젝트. 파일명에 남겨 어떤 건인지 알 수 있게."""
    parent = os.path.basename(os.path.dirname(rel)) if os.path.dirname(rel) else ""
    c = re.sub(r"\s+", "", parent)
    c = re.sub(r'[\\/:*?"<>|_]', "", c)   # 구분자 _ 와 충돌 방지
    return c[:22]


def _vision_classify(paths):
    """claude로 사진 단계 분류 → ['공정','완성',...] (실패 시 None들)."""
    listing = "\n".join(f"{i+1}. {p}" for i, p in enumerate(paths))
    prompt = ("아래 조형물 사진들을 Read 도구로 하나씩 열어보고 제작 단계를 분류하라.\n"
              "단계는 반드시 이 중 하나: 대표 / 공정 / 도장 / 설치 / 완성\n"
              "- 공정: 흰색·미완성·제작 중 (스티로폼/FRP 원형, 샌딩 등)\n"
              "- 도장: 채색 작업 중이거나 막 칠한 상태\n"
              "- 설치: 야외 현장에 설치된 모습\n"
              "- 완성: 완성된 조형물(실내/스튜디오 포함)\n"
              "- 대표: 가장 잘 나온 완성 대표컷\n\n"
              f"파일:\n{listing}\n\n"
              "오직 JSON 배열로만 출력(설명 금지): [{\"n\":1,\"stage\":\"공정\"}, ...]")
    out, err = run_claude(prompt, "haiku")    # 분류는 빠른 모델로
    if err or not out:
        return [None] * len(paths)
    m = re.search(r"\[.*\]", out, re.S)
    res = [None] * len(paths)
    if m:
        try:
            for it in json.loads(m.group(0)):
                i = int(it.get("n", 0)) - 1
                if 0 <= i < len(paths):
                    res[i] = it.get("stage")
        except Exception:
            pass
    return res


def _vision_folder_type(paths, type_words=None):
    """폴더(같은 프로젝트) 샘플 사진들을 보고 '종류' 한 단어로 판단 → 타입 문자열(실패 시 '')."""
    tw = type_words if type_words is not None else _TYPE_WORDS
    if not tw:
        return ""
    listing = "\n".join(f"{i+1}. {p}" for i, p in enumerate(paths))
    types = " / ".join(tw)
    prompt = ("아래 사진들은 같은 프로젝트의 조형물이다. Read 도구로 열어보고 이 프로젝트의 '종류'를 하나로 판단하라.\n"
              f"종류는 반드시 이 중 가장 맞는 하나: {types}\n"
              "- 캐릭터: 동물·사람형 귀여운 캐릭터 조형물\n"
              "- 마스코트: 기관·브랜드 마스코트 캐릭터\n"
              "- 글씨/스카시/간판: 글자·로고·간판 형태\n"
              "- 동상: 인물 동상  - 피규어: 작은 피규어  - 포토존: 포토존·벤치 구조물\n"
              "- 금속: 금속 재질  - FRP/EPS/스티로폼: 해당 재질 원형\n\n"
              f"파일:\n{listing}\n\n오직 위 목록의 단어 하나만 출력(설명 금지).")
    out, err = run_claude(prompt, "haiku")
    if err or not out:
        return ""
    for t in tw:
        if t.lower() in out.lower():
            return t
    return ""


def organize_worker(folder, type_words=None, stages=None):
    from concurrent.futures import ThreadPoolExecutor, as_completed
    from collections import defaultdict
    tw = type_words if type_words is not None else _TYPE_WORDS
    st = stages if stages is not None else _STAGE_HINTS
    stage_names = sorted(set(v for _, v in st), key=len, reverse=True) or ["완성"]
    done_re = re.compile("(" + "|".join(re.escape(s) for s in stage_names) + r")_\d{2}")
    try:
        files = list_images(folder, cap=10000)
        ORG.update(running=True, total=len(files), done=0, renamed=0, log=[], finished=False)
        # 1) 이미 정리된 것 제외
        todo = []
        for rel in files:
            name = os.path.basename(rel)
            ORG["done"] += 1
            if done_re.search(name):   # 이미 _단계_NN 형태
                continue
            todo.append(rel)
        ORG.update(total=len(todo), done=0)
        if not todo:
            ORG["log"].append("정리할 새 사진이 없습니다(이미 정리됨)."); return

        # 2) 폴더별 타입 결정 — 폴더명에 타입 있으면 그걸로, 없으면 비전 샘플로 1회
        byfolder = defaultdict(list)
        for rel in todo:
            byfolder[os.path.dirname(rel)].append(rel)
        folder_type = {d: _type_from_path((d or "") + "/_", tw) for d in byfolder}
        typeless = [d for d in byfolder if not folder_type[d]]
        if CLAUDE and tw and typeless:
            ORG["log"].append(f"폴더 타입 분류 ({len(typeless)}개 폴더, 비전)")

            def ftype(d):
                sample = [os.path.join(folder, r) for r in byfolder[d][:3]]
                return d, _vision_folder_type(sample, tw)
            with ThreadPoolExecutor(max_workers=4) as ex:
                for fut in as_completed([ex.submit(ftype, d) for d in typeless]):
                    try:
                        d, t = fut.result(); folder_type[d] = t
                    except Exception:
                        pass

        # 3) 단계 결정 — 이름에 힌트 있으면 그걸로, 없으면 비전(타입은 폴더값)
        plan, unclear = [], []
        for rel in todo:
            name = os.path.basename(rel)
            typ = _type_from_path(rel, tw) or folder_type.get(os.path.dirname(rel), "")
            stg = _stage_from_name(name, st)
            if stg:
                plan.append((rel, typ, stg)); ORG["done"] += 1
            else:
                unclear.append((rel, typ))
        if CLAUDE and unclear:
            B = 8
            batches = [unclear[i:i + B] for i in range(0, len(unclear), B)]
            ORG["log"].append(f"단계 분류 시작 ({len(unclear)}장 · {len(batches)}배치 동시)")

            def do_batch(batch):
                paths = [os.path.join(folder, r) for r, _ in batch]
                return list(zip(batch, _vision_classify(paths)))

            dn = 0
            with ThreadPoolExecutor(max_workers=4) as ex:
                for fut in as_completed([ex.submit(do_batch, b) for b in batches]):
                    try:
                        for (rel, typ), stg in fut.result():
                            plan.append((rel, typ, stg or "완성")); ORG["done"] += 1
                    except Exception:
                        pass
                    dn += 1
                    ORG["log"].append(f"단계 분류 {dn}/{len(batches)} 배치 완료")

        # 4) 이름변경 — [캠페인]_[타입]_[단계]_[번호]  (타입이 캠페인에 이미 있으면 생략)
        counters = {}
        for rel, typ, stg in plan:
            src = os.path.join(folder, rel)
            if not os.path.isfile(src):
                continue
            d = os.path.dirname(src)
            camp = _campaign_from_path(rel)
            parts = [p for p in [camp] if p]
            if typ and typ.lower() not in camp.lower():   # 캠페인에 없을 때만 타입 추가
                parts.append(typ)
            parts.append(stg)
            stem = "_".join(parts)
            key = (d, stem)
            counters[key] = counters.get(key, 0) + 1
            ext = os.path.splitext(src)[1]
            base = f"{stem}_{counters[key]:02d}"
            dst = os.path.join(d, base + ext)
            k = 1
            while os.path.exists(dst) and os.path.abspath(dst) != os.path.abspath(src):
                dst = os.path.join(d, f"{base}_{k}{ext}"); k += 1
            try:
                if os.path.abspath(dst) != os.path.abspath(src):
                    os.rename(src, dst); ORG["renamed"] += 1
            except Exception:
                pass
        ORG["log"].append(f"완료: {ORG['renamed']}장 이름변경 (타입+단계 태깅)")
    except Exception as e:
        ORG["log"].append("오류: " + str(e)[:150])
    finally:
        ORG.update(running=False, finished=True)


@app.route("/api/organize", methods=["POST"])
def api_organize():
    if ORG["running"]:
        return jsonify(ok=False, msg="이미 정리 중입니다.")
    body = request.get_json(force=True) or {}
    folder = body.get("folder", "")
    if not (folder and os.path.isdir(folder)):
        return jsonify(ok=False, msg="폴더를 먼저 지정하세요.")
    brand = brands.load_brand(body.get("brand") or "haofactory")
    threading.Thread(target=organize_worker, args=(folder, brand["type_words"], brand["stages"]),
                     daemon=True).start()
    return jsonify(ok=True)


@app.route("/api/organize-status")
def api_organize_status():
    return jsonify(running=ORG["running"], total=ORG["total"], done=ORG["done"],
                   renamed=ORG["renamed"], finished=ORG["finished"], log=ORG["log"][-6:])


@app.route("/img")
def img():
    folder = request.args.get("folder", ""); name = request.args.get("name", "")
    p = os.path.join(folder, name)
    if folder and os.path.isfile(p) and os.path.abspath(p).startswith(os.path.abspath(folder)):
        r = send_file(p)
        r.headers["Cache-Control"] = "no-store, max-age=0"
        return r
    return ("", 404)


import math, urllib.request
_BRAND_CORE = {"url": None, "data": None}


def _norm(s):
    return re.sub(r"\s+", "", str(s)).lower()


def fetch_brand_core(url):
    """홈페이지 title·meta keywords·h1에서 핵심 키워드/토큰 추출(메모리 캐시)."""
    if _BRAND_CORE["url"] == url and _BRAND_CORE["data"] is not None:
        return _BRAND_CORE["data"]
    kws, toks = [], set()
    if url:
        u = url if url.startswith("http") else "https://" + url
        try:
            html = urllib.request.urlopen(urllib.request.Request(u, headers={"User-Agent": "Mozilla/5.0"}), timeout=8).read().decode("utf-8", "ignore")
            def g(p):
                m = re.search(p, html, re.S | re.I); return m.group(1) if m else ""
            mk = g(r'name=["\']keywords["\'][^>]*content=["\']([^"\']+)')
            blob = re.sub(r"<[^>]+>", " ", " ".join([g(r"<title[^>]*>(.*?)</title>"), " ".join(re.findall(r"<h1[^>]*>(.*?)</h1>", html, re.S | re.I)), g(r'name=["\']description["\'][^>]*content=["\']([^"\']+)'), mk]))
            kws = [k.strip() for k in re.split(r"[,\n·|]", mk) if k.strip()]
            if len(kws) <= 1 and mk.strip():
                kws = [w for w in re.split(r"\s+", mk.strip()) if len(w) >= 2]
            for w in re.split(r"[^0-9A-Za-z가-힣]+", blob):
                if len(w) >= 2:
                    toks.add(w)
        except Exception:
            pass
    _BRAND_CORE["url"] = url
    _BRAND_CORE["data"] = {"norms": [_norm(k) for k in kws], "tok": sorted(toks)}
    return _BRAND_CORE["data"]


def _brand_fit(kw, core):
    norms, toks = core["norms"], core["tok"]
    if not norms and not toks:
        return 1.0
    nk = _norm(kw)
    if nk in norms:
        return 3.2
    if any(nk in c or c in nk for c in norms):
        return 2.4
    hits = sum(1 for t in toks if t and t in kw)
    return 1.9 if hits >= 2 else 1.1 if hits == 1 else 0.4


# 통검체크 Supabase(측정결과 공유) — 통검앱과 동일 anon 키
SUPA_URL = "https://stvyxslqkenupegjlste.supabase.co"
SUPA_KEY = ("eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.eyJpc3MiOiJzdXBhYmFzZSIsInJlZiI6InN0dnl4"
            "c2xxa2VudXBlZ2psc3RlIiwicm9sZSI6ImFub24iLCJpYXQiOjE3ODE5MTM1NzYsImV4cCI6MjA5"
            "NzQ4OTU3Nn0.8XXzhnRb_w26VFTfNsElsQJ9OfvYnmFdUOyu30bs3f4")


def supa_measurements(brand_name):
    """Supabase measurements_latest에서 이 브랜드 측정결과(노출/미노출) 가져오기 — 어느 PC든 공유."""
    if not brand_name:
        return {}
    import urllib.parse
    q = ("/rest/v1/measurements_latest?brand=eq.%s&select=keyword,summary,hits,measured_at"
         % urllib.parse.quote(brand_name))
    try:
        req = urllib.request.Request(SUPA_URL + q,
                                     headers={"apikey": SUPA_KEY, "Authorization": "Bearer " + SUPA_KEY})
        with urllib.request.urlopen(req, timeout=10) as r:
            rows = json.loads(r.read().decode() or "[]")
        return {row["keyword"]: {"summary": row.get("summary", ""), "hits": row.get("hits") or [],
                                 "date": (row.get("measured_at") or "")[:10]} for row in rows}
    except Exception:
        return {}


@app.route("/api/recommend")
def api_recommend():
    """통검 글감추천 — 키워드는 동봉 스냅샷(어느 PC든), 측정결과는 Supabase 공유."""
    brand = brands.load_brand(request.args.get("brand") or "haofactory")
    # 키워드 데이터: 동봉(brands/<id>/tonggeom_data.json) 우선 → 어느 PC에서도 작동. 없으면 로컬 통검 폴더.
    data = jload(os.path.join(brands.brand_dir(brand["id"]), "tonggeom_data.json"), None)
    if not data:
        tg = brand.get("tonggeom") or load_cfg()["tonggeom_dir"]
        data = jload(os.path.join(tg, "통검_데이터.json"), [])
    if not data:
        return jsonify(items=[], msg="이 브랜드의 통검 키워드 데이터가 없습니다. 프로그램을 껐다 켜서 최신 버전으로 업데이트하거나, 통검체크에서 측정을 한 번 돌려주세요.")
    # 측정결과: Supabase 공유(실시간) + 로컬 병합(로컬이 더 최신이면)
    meas = supa_measurements(brand.get("name") or "")
    tg = brand.get("tonggeom") or load_cfg()["tonggeom_dir"]
    for k, v in jload(os.path.join(tg, "통검_측정.json"), {}).items():
        if k not in meas or v.get("date", "") > meas[k].get("date", ""):
            meas[k] = v
    core = fetch_brand_core(brand.get("homepage") or "")
    impw = {"상": 3.0, "중": 2.0, "하": 1.0, "x": 0.0, "": 1.5}
    out = []
    for r in data:
        iw = impw.get(str(r.get("importance", "")).strip(), 1.5)
        if iw <= 0:
            continue
        vol = r.get("volume") or 0
        m = meas.get(r["keyword"])
        state = "미측정" if not m else ("노출" if m.get("hits") else "미노출")
        opp = 0.6 if state == "노출" else 1.0
        bare = r["keyword"].replace(" ", "")
        ach = 0.3 if (len(bare) <= 3 and vol >= 2000) else 1.0
        fit = _brand_fit(r["keyword"], core)
        sc = (fit ** 1.6) * iw * math.log10(min(vol, 12000) + 10) * opp * ach
        out.append({"keyword": r["keyword"], "category": r.get("category", ""),
                    "volume": vol, "importance": r.get("importance", ""), "state": state,
                    "score": round(sc, 2)})
    out.sort(key=lambda x: -x["score"])
    return jsonify(items=out[:40])


# ── 브랜드 관리 ──────────────────────────────────────────
@app.route("/api/version")
def api_version():
    return jsonify(version=jload(os.path.join(HERE, "version.json"), {}).get("version", "?"))


@app.route("/api/brands")
def api_brands():
    return jsonify(brands=brands.list_brands())


@app.route("/api/brand")
def api_brand():
    return jsonify(brands.load_brand(request.args.get("id") or "haofactory"))


@app.route("/api/brand-save", methods=["POST"])
def api_brand_save():
    cfg = request.get_json(force=True) or {}
    if not (cfg.get("name") or "").strip():
        return jsonify(ok=False, msg="브랜드 이름을 입력하세요.")
    # 콤마 문자열로 들어온 type_words 정규화
    tw = cfg.get("type_words")
    if isinstance(tw, str):
        cfg["type_words"] = [t.strip() for t in re.split(r"[,\n]", tw) if t.strip()]
    ch = cfg.get("card_headlines")
    if isinstance(ch, str):
        cfg["card_headlines"] = [t.strip() for t in ch.split("\n") if t.strip()]
    bid = brands.save_brand(cfg)
    return jsonify(ok=True, id=bid, brand=brands.load_brand(bid))


# 통검체크 dist 루트 + 브랜드 id/색 매핑
TONGGEOM_DIST = r"C:/Users/이건희/trevity-노출체크/dist"
_TG_MAP = {  # 통검 폴더 브랜드명 → (id, 색)
    "하오팩토리": ("haofactory", "#FD6F22"), "하오디자인": ("haodesign", "#16A34A"),
    "하오스튜디오": ("haostudio", "#7C3AED"), "퍼스트디자인": ("firstdesign", "#3B82F6"),
    "레드트랜스": ("redtrans", "#E11D48"), "윈차이나": ("winchina", "#DC2626"),
}


def sync_tonggeom_brands():
    """통검체크 dist의 '<브랜드> 통검체크' 폴더를 스캔 → 브랜드 생성/연결(기존 커스텀 필드는 보존)."""
    added, linked = [], []
    if not os.path.isdir(TONGGEOM_DIST):
        return {"added": added, "linked": linked, "err": "dist 폴더 없음"}
    for d in sorted(os.listdir(TONGGEOM_DIST)):
        full = os.path.join(TONGGEOM_DIST, d)
        if not (os.path.isdir(full) and d.endswith("통검체크")):
            continue
        bname = d.replace("통검체크", "").strip()
        bid, color = _TG_MAP.get(bname, (brands.slugify(bname), "#3B82F6"))
        hp = ""
        try:
            hp = jload(os.path.join(full, "통검_설정.json"), {}).get("homepage", "")
        except Exception:
            pass
        existed = os.path.isfile(os.path.join(brands.brand_dir(bid), "brand.json"))
        cfg = brands.load_brand(bid) if existed else dict(brands.DEFAULTS)
        cfg["id"] = bid
        cfg["name"] = bname
        cfg["tonggeom"] = full
        if hp:
            cfg["homepage"] = hp if hp.startswith("http") else "https://" + hp
        if not existed:
            cfg["color"] = color
        brands.save_brand(cfg)
        # 키워드 목록(통검_데이터.json)을 브랜드 폴더에 동봉 → 배포 시 어느 PC에서도 추천 작동
        try:
            kw = os.path.join(full, "통검_데이터.json")
            if os.path.isfile(kw):
                shutil.copy(kw, os.path.join(brands.brand_dir(bid), "tonggeom_data.json"))
        except Exception:
            pass
        (linked if existed else added).append(bname)
    return {"added": added, "linked": linked}


@app.route("/api/sync-brands", methods=["POST"])
def api_sync_brands():
    return jsonify(ok=True, **sync_tonggeom_brands())


@app.route("/api/preview-cards", methods=["POST"])
def api_preview_cards():
    """브랜드 카드 템플릿을 샘플 사진/문구로 렌더 — 최종 체크용 미리보기."""
    b = request.get_json(force=True) or {}
    bid = b.get("brand") or "haofactory"
    tpl = str(b.get("template") or "1")
    brand = brands.load_brand(bid)
    adir = brands.assets_dir(bid, tpl)
    if not os.path.isfile(os.path.join(adir, "layout.json")):
        return jsonify(ok=False, msg="이 브랜드는 카드 템플릿이 없습니다.")
    # 샘플 사진: 탭 폴더 → 데스크톱 drive-download → Pictures 순
    files, folder = [], b.get("folder") or ""
    if folder and os.path.isdir(folder):
        files = list_images(folder, cap=300)
    if not files:
        for cand in glob.glob(os.path.join(os.path.expanduser("~"), "Desktop", "drive-download*")) + \
                    [os.path.join(os.path.expanduser("~"), "Pictures")]:
            if os.path.isdir(cand):
                fl = list_images(cand, cap=300)
                if fl:
                    folder, files = cand, fl
                    break
    random.shuffle(files)
    n = len(jload(os.path.join(adir, "layout.json"), {}).get("slides", [])) or 7
    photo_paths = [os.path.join(folder, f) for f in files[:n]] or [""]
    # 샘플 문구
    heads_raw = brand.get("card_headlines") or []
    heads = []
    for i in range(7):
        h = re.sub(r"\([^)]*\)", "", heads_raw[i]).strip() if i < len(heads_raw) else ""
        heads.append(h or f"샘플 헤드라인 {i+1}")
    bodies = ["샘플 본문 설명이 들어가는 자리입니다."] * 7
    pdir = os.path.join(OUT_DIR, "_preview_" + bid)
    shutil.rmtree(pdir, ignore_errors=True)
    try:
        res = cardnews_pil.make_cards(photo_paths, heads, pdir, adir,
                                      subtitle="이 글 주제를 한 줄로 요약한 부제",
                                      bodies=bodies, title=f"{brand['name']} 샘플 제목, 핵심을 한 줄로 보여주는 미리보기")
        return jsonify(ok=True, dir=pdir, pngs=res["pngs"], name=brand["name"])
    except Exception as e:
        return jsonify(ok=False, msg=str(e)[:200])


def _sample_photo():
    """썸네일용 샘플 사진 1장(데스크톱 drive-download → Pictures)."""
    for cand in glob.glob(os.path.join(os.path.expanduser("~"), "Desktop", "drive-download*")) + \
                [os.path.join(os.path.expanduser("~"), "Pictures")]:
        if os.path.isdir(cand):
            fl = list_images(cand, cap=50)
            if fl:
                return os.path.join(cand, fl[0])
    return ""


@app.route("/api/template-thumb")
def api_template_thumb():
    """카드 디자인 슬롯의 표지를 썸네일로 렌더(선택기용). layout.json 바뀌면 자동 갱신."""
    bid = request.args.get("brand", "haofactory")
    tpl = str(request.args.get("template", "1") or "1")
    adir = brands.assets_dir(bid, tpl)
    lp = os.path.join(adir, "layout.json")
    if not os.path.isfile(lp):
        return ("", 404)
    cache = os.path.join(OUT_DIR, f"_thumb_{bid}_{tpl}.png")
    if (not os.path.isfile(cache)) or os.path.getmtime(cache) < os.path.getmtime(lp):
        brand = brands.load_brand(bid)
        heads = brand.get("card_headlines") or []
        kw = (re.sub(r"\([^)]*\)", "", heads[0]).strip() if heads else "") or "샘플 키워드"
        try:
            cardnews_pil.render_card(0, _sample_photo(), kw, (210, 90, 120), cache, adir,
                                     subtitle="주제를 한 줄로 요약한 부제", title=kw)
        except Exception as e:
            return (str(e)[:120], 500)
    r = send_file(cache)
    r.headers["Cache-Control"] = "no-store, max-age=0"
    return r


@app.route("/api/brand-template", methods=["POST"])
def api_brand_template():
    """브랜드 카드뉴스 PPTX 템플릿 등록 → 장식/레이아웃 추출."""
    b = request.get_json(force=True) or {}
    bid = b.get("id") or ""
    pptx = b.get("pptx") or ""
    if not (bid and os.path.isfile(pptx) and pptx.lower().endswith(".pptx")):
        return jsonify(ok=False, msg="브랜드 또는 PPTX 경로가 올바르지 않습니다.")
    try:
        r = cardnews_pil.extract_template(pptx, brands.assets_dir(bid))
        return jsonify(ok=True, **r)
    except Exception as e:
        return jsonify(ok=False, msg=str(e)[:200])


@app.route("/api/brand-template-file", methods=["POST"])
def api_brand_template_file():
    """브랜드 카드뉴스 PPTX 직접 업로드 → 추출(파일 선택 방식)."""
    bid = request.form.get("id", "")
    slot = str(request.form.get("slot", "1") or "1")
    f = request.files.get("file")
    if not (bid and f and f.filename.lower().endswith(".pptx")):
        return jsonify(ok=False, msg="브랜드와 PPTX 파일이 필요합니다.")
    tmp = os.path.join(OUT_DIR, f"_tpl_{bid}.pptx")
    try:
        f.save(tmp)
        r = cardnews_pil.extract_template(tmp, brands.assets_dir(bid, slot))
        return jsonify(ok=True, slot=slot, **r)
    except Exception as e:
        return jsonify(ok=False, msg=str(e)[:200])
    finally:
        try:
            os.remove(tmp)
        except Exception:
            pass


@app.route("/api/generate", methods=["POST"])
def api_generate():
    b = request.get_json(force=True)
    keyword = (b.get("keyword") or "").strip()
    folder = b.get("folder") or ""
    hint = (b.get("hint") or "").strip()
    tpl = str(b.get("template") or "1")   # 카드 디자인 슬롯(b는 아래 루프에서 덮어써지므로 먼저 읽음)
    if not keyword:
        return jsonify(ok=False, msg="키워드를 입력하세요.")
    brand = brands.load_brand(b.get("brand") or "haofactory")
    files = list_images(folder, cap=20000)              # 전체 다 가져와서
    files = filter_by_keyword(files, keyword, brand["type_words"])   # 키워드 매칭(브랜드 타입단어)
    random.shuffle(files)                               # ★매 생성마다 다른 순서(Claude도 매번 다르게 고름)
    files = files[:150]
    model = (b.get("model") or load_cfg().get("model") or "opus")
    prompt = build_prompt(keyword, files, hint, brand)
    out, err = run_claude(prompt, model)
    if err:
        return jsonify(ok=False, msg=err)
    # 카드뉴스 헤드라인 분리
    parts = re.split(r"\n\s*카드뉴스\s*[:：]\s*\n?", out, maxsplit=1)
    cards, bodies = [], []
    if len(parts) > 1:
        for ln in parts[1].split("\n"):
            mm = re.match(r"^\s*\d+\s*[.)]\s*(.+)$", ln.strip())
            if mm:
                raw = mm.group(1)
                hpart, bpart = (raw.split("|", 1) + [""])[:2] if "|" in raw else (raw, "")
                h = re.sub(r"\([^)]*\)", "", hpart).strip() or hpart.strip()
                b = re.sub(r"\([^)]*\)", "", bpart).strip()
                cards.append(h); bodies.append(b)
    sub = ""
    ms = re.search(r"표지\s*부제\s*[:：]\s*(.+)", out)
    if ms:
        sub = re.sub(r"\([^)]*\)", "", ms.group(1)).strip() or ms.group(1).strip()
    post = parse_manuscript(parts[0], files)
    post["folder"] = folder
    post["keyword"] = keyword
    post["cards"] = cards[:7]
    post["card_bodies"] = bodies[:7]
    post["subtitle"] = sub
    post["raw"] = out
    # 카드뉴스 생성 (순수 PIL)
    post["cardnews"] = ""
    # 카드 7장 = 서로 다른 사진(Claude 매칭 우선 + 폴더 전체에서 분산 채움)
    claude_picks = [b["file"] for b in post["blocks"] if b["type"] == "photo" and b.get("file")]
    adir = brands.assets_dir(brand["id"], tpl)
    has_tpl = os.path.isfile(os.path.join(adir, "layout.json"))
    n_cards = 7
    if has_tpl:
        try:
            n_cards = len(jload(os.path.join(adir, "layout.json"), {}).get("slides", [])) or 7
        except Exception:
            n_cards = 7
    card_files = pick_diverse(files, claude_picks, n_cards)
    post["cardnews_pngs"] = []
    post["cardnews_job"] = ""
    if card_files and len(cards) >= 6 and has_tpl:
        import uuid
        jid = uuid.uuid4().hex[:10]
        CARD_JOBS[jid] = {"status": "rendering"}
        post["cardnews_job"] = jid
        photo_paths = [os.path.join(folder, f) for f in card_files]
        threading.Thread(target=_cardnews_job, args=(jid, photo_paths, cards, keyword, adir, sub, bodies, post.get("title", "")), daemon=True).start()
    return jsonify(ok=True, post=post)


@app.route("/api/cardnews-status")
def api_cardnews_status():
    j = CARD_JOBS.get(request.args.get("id", ""), {"status": "unknown"})
    return jsonify(j)


@app.route("/api/swap-card", methods=["POST"])
def api_swap_card():
    """카드 1장의 사진만 교체 → 그 카드만 재렌더(전체 재생성 X)."""
    b = request.get_json(force=True)
    cdir = b.get("cardnews_dir", "")
    folder = b.get("folder", "")
    idx = int(b.get("index", 0))
    photo = os.path.join(folder, b.get("file", ""))
    if not (cdir and os.path.isdir(cdir) and os.path.isfile(photo)):
        return jsonify(ok=False, msg="파일을 찾을 수 없습니다.")
    try:
        png = cardnews_pil.edit_card(cdir, idx, src=photo)
        return jsonify(ok=True, png=png, name=os.path.basename(png), src=photo)
    except Exception as e:
        return jsonify(ok=False, msg=str(e)[:200])


@app.route("/api/card-text", methods=["POST"])
def api_card_text():
    """카드 글(헤드라인/부제/본문/제목) 수정 → 그 카드만 재렌더."""
    b = request.get_json(force=True)
    cdir = b.get("cardnews_dir", "")
    idx = int(b.get("index", 0))
    if not (cdir and os.path.isdir(cdir)):
        return jsonify(ok=False, msg="카드 폴더를 찾을 수 없습니다.")
    kw = {k: b[k] for k in ("headline", "subtitle", "body", "title") if k in b}
    try:
        png = cardnews_pil.edit_card(cdir, idx, **kw)
        return jsonify(ok=True, png=png, name=os.path.basename(png))
    except Exception as e:
        return jsonify(ok=False, msg=str(e)[:200])


@app.route("/api/adjust-card", methods=["POST"])
def api_adjust_card():
    """같은 사진을 위치(cx,cy)·확대(zoom)만 바꿔 그 카드만 재렌더."""
    b = request.get_json(force=True)
    cdir = b.get("cardnews_dir", "")
    idx = int(b.get("index", 0))
    cx = float(b.get("cx", 0.5)); cy = float(b.get("cy", 0.5)); zoom = float(b.get("zoom", 1.0))
    if not (cdir and os.path.isdir(cdir)):
        return jsonify(ok=False, msg="카드 폴더를 찾을 수 없습니다.")
    try:
        png = cardnews_pil.edit_card(cdir, idx, center=(cx, cy), zoom=zoom)
        return jsonify(ok=True, png=png, name=os.path.basename(png))
    except Exception as e:
        return jsonify(ok=False, msg=str(e)[:200])


@app.route("/api/open", methods=["POST"])
def api_open():
    p = (request.get_json(force=True) or {}).get("path", "")
    if p and os.path.exists(p):
        try:
            os.startfile(p)
        except Exception:
            pass
        return jsonify(ok=True)
    return jsonify(ok=False)


@app.route("/api/docx", methods=["POST"])
def api_docx():
    from docx import Document
    from docx.shared import Cm, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    b = request.get_json(force=True)
    post = b.get("post", {}); folder = post.get("folder", "")
    FONT = "Pretendard"

    def setf(run, sz, bold=False):
        run.font.name = FONT; run.font.size = Pt(sz); run.font.bold = bold
        run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)

    pngs = post.get("cardnews_pngs", [])
    use_cards = bool(pngs) and post.get("use_cards", True)
    pidx = [0]
    doc = Document()
    for s in doc.sections:
        s.left_margin = s.right_margin = Cm(2.3)
    if post.get("title"):
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(14)
        setf(p.add_run(post["title"]), 16, True)
    for blk in post.get("blocks", []):
        if blk["type"] == "photo":
            img = None
            if use_cards:
                img = pngs[pidx[0]] if pidx[0] < len(pngs) else None   # 카드 다 쓰면 여분 자리 생략
            elif blk.get("file"):
                img = os.path.join(folder, blk["file"])
            pidx[0] += 1
            if img and os.path.isfile(img):
                pp = doc.add_paragraph(); pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                pp.paragraph_format.space_before = Pt(6); pp.paragraph_format.space_after = Pt(12)
                try:
                    pp.add_run().add_picture(img, width=Cm(14))
                except Exception:
                    pass
        elif blk["type"] == "text":
            for para in blk["text"].split("\n\n"):
                lines = [x for x in para.split("\n") if x.strip()]
                if not lines:
                    continue
                hm = re.match(r"^\*\*(.+?)\*\*$", lines[0].strip())
                pp = doc.add_paragraph(); pp.paragraph_format.line_spacing = 1.5
                pp.paragraph_format.space_after = Pt(9)
                if hm and len(lines) == 1:
                    setf(pp.add_run(hm.group(1)), 13, True)
                else:
                    first = True
                    for ln in lines:
                        if not first:
                            pp.add_run().add_break()
                        setf(pp.add_run(re.sub(r"\*\*", "", ln)), 11)
                        first = False
    # 본문 자리보다 카드가 많으면(예: 자리 6개·카드 7장) 남는 카드는 끝에 붙임
    if use_cards:
        while pidx[0] < len(pngs):
            img = pngs[pidx[0]]; pidx[0] += 1
            if os.path.isfile(img):
                pp = doc.add_paragraph(); pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                pp.paragraph_format.space_before = Pt(6); pp.paragraph_format.space_after = Pt(12)
                try:
                    pp.add_run().add_picture(img, width=Cm(14))
                except Exception:
                    pass
    safe = re.sub(r"[\\/:*?\"<>|]", "_", post.get("keyword", "글"))
    stamp = datetime.datetime.now().strftime("%H%M%S")
    out = os.path.join(OUT_DIR, f"{safe}_{stamp}.docx")
    doc.save(out)
    try:
        os.startfile(out)
    except Exception:
        pass
    return jsonify(ok=True, path=out)


@app.after_request
def _no_cache(resp):
    # 브라우저가 옛 화면(HTML/JS)을 캐시해 새 기능이 안 보이는 문제 방지
    resp.headers["Cache-Control"] = "no-store, no-cache, must-revalidate, max-age=0"
    resp.headers["Pragma"] = "no-cache"
    resp.headers["Expires"] = "0"
    return resp


@app.route("/")
def index():
    return PAGE


PAGE = r"""<!DOCTYPE html><html lang="ko"><head><meta charset="utf-8">
<meta name="viewport" content="width=device-width,initial-scale=1"><title>블로그 작성기</title>
<style>
:root{--brand:#FD6F22;--brand-d:#e85f12;--brand-l:#fff3ec;--ink:#1b2330;--ink2:#46506180;
  --muted:#8b95a5;--line:#e9ecf1;--line2:#f1f3f7;--bg:#f3f5f8;--card:#fff;
  --ok:#079455;--okbg:#e7f7ee;--warn:#dc8400;--warnbg:#fff3df;--err:#e0483c;--errbg:#fdecea;
  --sh:0 1px 2px rgba(16,24,40,.04),0 6px 20px rgba(16,24,40,.05);
  --sh-lg:0 16px 48px rgba(16,24,40,.18);--r:16px}
*{box-sizing:border-box;font-family:'Pretendard','Malgun Gothic',sans-serif}
body{margin:0;background:var(--bg);color:var(--ink);-webkit-font-smoothing:antialiased;font-size:14px}
::selection{background:var(--brand-l)}
.top{background:rgba(255,255,255,.85);backdrop-filter:blur(10px);border-bottom:1px solid var(--line);
  padding:14px 26px;display:flex;align-items:center;gap:13px;position:sticky;top:0;z-index:20}
.top .logo{width:30px;height:30px;border-radius:9px;background:linear-gradient(135deg,var(--brand),var(--brand-2));
  display:flex;align-items:center;justify-content:center;font-size:16px;box-shadow:0 4px 12px var(--brand-sh)}
.top h1{font-size:16px;margin:0;font-weight:800;letter-spacing:-.2px}
.top .st{font-size:12px;color:var(--muted);font-weight:600}
.pill{margin-left:auto;font-size:11.5px;font-weight:700;padding:5px 12px;border-radius:99px;
  background:var(--okbg);color:var(--ok);display:inline-flex;align-items:center;gap:5px}
.pill.off{background:var(--warnbg);color:var(--warn)}
.tabs{display:flex;gap:5px;align-items:center;padding:14px 26px 0;flex-wrap:wrap;max-width:1120px;margin:0 auto}
.tab{padding:9px 16px;border:1px solid transparent;border-radius:11px 11px 0 0;background:transparent;
  cursor:pointer;font-size:13px;font-weight:700;color:var(--muted);transition:.15s;position:relative;top:1px}
.tab:hover{color:var(--ink2);background:#fff8}
.tab.on{background:var(--card);color:var(--ink);border-color:var(--line);border-bottom-color:var(--card);box-shadow:0 -2px 8px rgba(16,24,40,.03)}
.tab.on::before{content:"";position:absolute;left:14px;right:14px;top:0;height:2.5px;border-radius:2px;background:var(--brand)}
.tab .x{margin-left:8px;color:#c2c8d2;font-weight:800;border-radius:5px;padding:0 3px}
.tab .x:hover{color:var(--err);background:#fff}
.addtab{padding:8px 13px;border:1px dashed var(--line);border-radius:9px;background:#fff;cursor:pointer;font-size:13px;color:var(--muted);font-weight:700;transition:.15s}
.addtab:hover{color:var(--brand);border-color:var(--brand)}
.wrap{max-width:1120px;margin:0 auto;padding:0 26px 80px}
.card{background:var(--card);border:1px solid var(--line);border-radius:0 var(--r) var(--r) var(--r);padding:24px 26px;box-shadow:var(--sh)}
.field{margin-bottom:18px}
.field label{display:block;font-size:12px;font-weight:800;color:var(--ink2);margin-bottom:7px;letter-spacing:.1px}
.row{display:flex;gap:8px;align-items:center}
input[type=text],textarea{width:100%;border:1.5px solid var(--line);border-radius:11px;padding:11px 13px;font-size:13.5px;color:var(--ink);background:#fdfdfe;transition:.15s;outline:none}
input[type=text]:focus,textarea:focus{border-color:var(--brand);background:#fff;box-shadow:0 0 0 3.5px var(--brand-l)}
input[type=text]::placeholder,textarea::placeholder{color:#b4bcc8}
textarea{height:58px;resize:vertical;line-height:1.5}
select{border:1.5px solid var(--line);border-radius:10px;padding:10px 11px;font-size:12.5px;font-weight:700;color:var(--ink2);background:#fdfdfe;cursor:pointer;outline:none}
select:focus{border-color:var(--brand)}
.btn{padding:11px 17px;border-radius:11px;border:1.5px solid var(--line);background:#fff;font-weight:700;font-size:13px;cursor:pointer;white-space:nowrap;transition:.15s;color:var(--ink2)}
.btn:hover{background:#f7f8fa;border-color:#dde1e8;color:var(--ink)}
.btn:active{transform:translateY(1px)}
.btn.pri{background:linear-gradient(135deg,var(--brand),var(--brand-2));border-color:transparent;color:#fff;box-shadow:0 4px 14px var(--brand-sh)}
.btn.pri:hover{filter:brightness(1.05);box-shadow:0 7px 20px var(--brand-sh);color:#fff}
.btn:disabled{opacity:.6;cursor:default;transform:none}
.btn.lg{padding:13px 26px;font-size:14px}
.photos{display:flex;gap:7px;flex-wrap:wrap;margin-top:10px}
.photos img{width:64px;height:64px;object-fit:cover;border-radius:9px;border:1px solid var(--line);transition:.15s}
.photos img:hover{transform:scale(1.06);box-shadow:var(--sh)}
.hint{margin-top:9px;font-size:12.5px;font-weight:700;display:flex;align-items:center;gap:6px}
.hint.ok{color:var(--ok)}.hint.bad{color:var(--err)}
.br-item{padding:11px 18px;font-size:13px;cursor:pointer;border-bottom:1px solid var(--line2);transition:.12s}
.br-item:hover{background:var(--brand-l);color:var(--brand-d)}
.recbox{border:1px solid var(--line);border-radius:12px;max-height:260px;overflow:auto;margin-top:10px;display:none;box-shadow:var(--sh)}
.recbox table{width:100%;border-collapse:collapse;font-size:12.5px}
.recbox td{padding:9px 12px;border-bottom:1px solid var(--line2);cursor:pointer}
.recbox tr:hover td{background:var(--brand-l)}
.recst{font-size:10px;font-weight:800;padding:3px 8px;border-radius:99px}
.recst.미노출{background:var(--warnbg);color:var(--warn)}.recst.노출{background:var(--okbg);color:var(--ok)}.recst.미측정{background:#eef1f4;color:#889}
.preview{margin-top:22px;border-top:1px solid var(--line);padding-top:20px;display:none}
.pv-head{display:flex;align-items:flex-start;gap:12px;flex-wrap:wrap;margin-bottom:6px}
.pv-title{font-size:19px;font-weight:800;letter-spacing:-.3px;flex:1;min-width:200px}
.badges{display:flex;gap:6px;flex-wrap:wrap;margin:10px 0 16px}
.badge{font-size:11.5px;font-weight:700;padding:5px 11px;border-radius:99px;background:#eef1f4;color:var(--ink2);display:inline-flex;align-items:center;gap:5px}
.badge.ok{background:var(--okbg);color:var(--ok)}.badge.warn{background:var(--warnbg);color:var(--warn)}.badge.brand{background:var(--brand-l);color:var(--brand-d)}
.pv-body p{white-space:pre-wrap;line-height:1.75;font-size:14px;margin:11px 0;color:#2b3340}
.pv-body h3{font-size:15.5px;font-weight:800;margin:22px 0 8px;padding-left:11px;border-left:3px solid var(--brand)}
.pv-body img{max-width:320px;border-radius:12px;display:block;margin:12px 0;box-shadow:var(--sh)}
.cardsec-hint{font-size:11.5px;color:var(--muted);font-weight:600;margin:6px 0 2px;display:flex;align-items:center;gap:6px}
.cardwrap{position:relative;display:inline-block;margin:10px 0;border-radius:14px;overflow:hidden;box-shadow:var(--sh);transition:.15s}
.cardwrap:hover{box-shadow:var(--sh-lg)}
.cardwrap img{display:block;border-radius:14px}
.cardbar{position:absolute;left:50%;bottom:10px;transform:translateX(-50%) translateY(6px);display:flex;align-items:center;gap:3px;
  background:rgba(20,24,32,.74);padding:5px 6px;border-radius:12px;opacity:0;transition:.16s;backdrop-filter:blur(6px);box-shadow:0 6px 18px rgba(0,0,0,.3)}
.cardwrap:hover .cardbar{opacity:1;transform:translateX(-50%) translateY(0)}
.cardbar button{height:28px;min-width:28px;border:0;border-radius:7px;background:rgba(255,255,255,.92);color:#222;font-size:12px;font-weight:800;cursor:pointer;line-height:1;padding:0 7px;transition:.1s}
.cardbar button.tb{font-size:11.5px}
.cardbar button:hover{background:var(--brand);color:#fff}
.cardbar .sep{width:1px;height:18px;background:rgba(255,255,255,.28);margin:0 2px}
#swaplist img{transition:.12s}
#swaplist img:hover{border-color:var(--brand)!important;transform:scale(1.06)}
.modal-bg{display:none;position:fixed;inset:0;background:rgba(16,22,34,.45);backdrop-filter:blur(2px);z-index:50}
.modal{background:#fff;border-radius:18px;overflow:hidden;box-shadow:var(--sh-lg)}
.pv-photo-missing{font-size:12px;color:var(--err);background:var(--errbg);padding:7px 12px;border-radius:9px;display:inline-block}
.spin{display:inline-block;width:14px;height:14px;border:2px solid currentColor;border-top-color:transparent;border-radius:50%;animation:sp .7s linear infinite;vertical-align:-2px}
@keyframes sp{to{transform:rotate(360deg)}}
.warn{background:var(--warnbg);border:1px solid #fcd9a3;color:#a85f00;padding:12px 16px;border-radius:12px;font-size:12.5px;margin:14px 26px;font-weight:600;max-width:1120px;margin-left:auto;margin-right:auto}
.gen-row{display:flex;gap:10px;align-items:center;background:var(--brand-l);border:1px solid transparent;border-radius:14px;padding:14px 16px;margin-top:4px;flex-wrap:wrap}
#toasts{position:fixed;right:22px;bottom:22px;z-index:90;display:flex;flex-direction:column;gap:9px}
.toast{background:#1f2733;color:#fff;padding:12px 16px;border-radius:12px;font-size:13px;font-weight:600;box-shadow:var(--sh-lg);display:flex;align-items:center;gap:9px;animation:tin .25s ease;max-width:340px}
.toast.ok{background:#0e7a48}.toast.err{background:#c33}.toast b{font-weight:800}
@keyframes tin{from{opacity:0;transform:translateY(10px)}to{opacity:1;transform:none}}
.spacer{flex:1}
.opt{margin-bottom:18px;border:1px solid var(--line);border-radius:11px;overflow:hidden}
.opt summary{cursor:pointer;font-size:12.5px;font-weight:700;color:var(--ink2);padding:11px 14px;list-style:none;user-select:none;background:#fafbfc}
.opt summary::-webkit-details-marker{display:none}
.opt summary:hover{color:var(--brand)}
.opt[open] summary{border-bottom:1px solid var(--line)}
.opt textarea{border:0;border-radius:0;width:100%}
.naver-box{background:var(--brand-l);border:1px solid transparent;border-radius:13px;padding:14px 16px;margin:4px 0}
.naver-title{font-weight:800;font-size:14px;margin-bottom:10px}
.naver-steps{font-size:12px;color:var(--ink2);margin-top:10px;line-height:1.6;font-weight:600}
.naver-steps b{color:var(--brand-d)}
.tplrow{display:flex;gap:10px;flex-wrap:wrap}
.tplopt{cursor:pointer;border:2.5px solid transparent;border-radius:13px;padding:5px;background:#f3f5f7;transition:.15s;text-align:center}
.tplopt:hover{background:#e9edf1}
.tplopt img{width:124px;height:124px;object-fit:cover;border-radius:9px;display:block;background:#fff}
.tplopt span{font-size:12px;font-weight:800;color:var(--ink2);display:block;margin-top:5px}
.tplopt.sel{border-color:var(--brand);background:var(--brand-l)}
.tplopt.sel span{color:var(--brand-d)}
</style></head><body>
<div class="top"><div class="logo">🖊</div><h1>블로그 작성기</h1>
  <select id="brandsel" onchange="switchBrand(this.value)" title="브랜드 선택"></select>
  <button class="btn" style="padding:7px 12px" onclick="openBrandMgr()">⚙ 브랜드</button>
  <span class="pill" id="engine">확인 중…</span>
  <button class="btn" style="padding:7px 12px" onclick="openAccounts()" title="Claude 계정 관리">👤 계정</button>
  <span id="ver" style="font-size:11px;color:var(--muted);font-weight:700"></span></div>
<div id="loginwarn" class="warn" style="display:none"></div>
<div id="toasts"></div>
<div class="tabs" id="tabs"></div>
<div class="wrap"><div class="card" id="card"></div></div>
<div id="brmodal" style="display:none;position:fixed;inset:0;background:rgba(0,0,0,.35);z-index:50" onclick="if(event.target==this)closeBrowse()">
  <div style="background:#fff;max-width:560px;margin:7vh auto;border-radius:14px;overflow:hidden;box-shadow:0 20px 60px rgba(0,0,0,.25)">
    <div style="padding:14px 18px;border-bottom:1px solid var(--line);display:flex;align-items:center;gap:10px">
      <b style="font-size:14px">📁 사진 폴더 선택</b><span style="flex:1"></span>
      <button class="btn" onclick="closeBrowse()">닫기</button></div>
    <div id="brplaces" style="padding:9px 18px;display:flex;gap:6px;flex-wrap:wrap;border-bottom:1px solid var(--line)"></div>
    <div style="padding:9px 18px;font-size:12px;color:var(--muted);word-break:break-all" id="brpath"></div>
    <div id="brlist" style="max-height:48vh;overflow:auto;border-top:1px solid var(--line)"></div>
    <div style="padding:12px 18px;border-top:1px solid var(--line);display:flex;align-items:center;gap:10px">
      <span id="brimg" style="font-size:12px;color:var(--muted)"></span><span style="flex:1"></span>
      <button class="btn pri" id="brpick" onclick="selectFolder()" disabled>✓ 이 폴더 선택</button></div>
  </div></div>
<div id="swapmodal" style="display:none;position:fixed;inset:0;background:rgba(0,0,0,.35);z-index:55" onclick="if(event.target==this)closeSwap()">
  <div style="background:#fff;max-width:640px;margin:8vh auto;border-radius:14px;overflow:hidden;box-shadow:0 20px 60px rgba(0,0,0,.25)">
    <div style="padding:14px 18px;border-bottom:1px solid var(--line);display:flex;align-items:center;gap:10px">
      <b style="font-size:14px">🔄 이 카드 사진 교체</b><span style="font-size:12px;color:var(--muted)">바꿀 사진 클릭</span><span style="flex:1"></span>
      <button class="btn" onclick="closeSwap()">닫기</button></div>
    <div id="swaplist" style="padding:14px 18px;display:flex;gap:7px;flex-wrap:wrap;max-height:62vh;overflow:auto"></div>
  </div></div>
<div id="textmodal" class="modal-bg" onclick="if(event.target==this)el('textmodal').style.display='none'">
  <div class="modal" style="max-width:520px;margin:9vh auto">
    <div style="padding:15px 20px;border-bottom:1px solid var(--line);display:flex;align-items:center;gap:10px">
      <b style="font-size:15px">✏️ 카드 글 수정</b><span class="spacer"></span>
      <button class="btn" onclick="el('textmodal').style.display='none'">닫기</button></div>
    <div id="txtfields" style="padding:18px 20px;max-height:60vh;overflow:auto"></div>
    <div style="padding:0 20px 18px;display:flex;gap:8px"><button class="btn pri" onclick="applyText()">적용</button>
      <span style="font-size:11.5px;color:var(--muted);align-self:center">수정 후 그 카드만 다시 그려집니다</span></div>
  </div></div>
<div id="previewmodal" class="modal-bg" style="z-index:60" onclick="if(event.target==this)el('previewmodal').style.display='none'">
  <div class="modal" style="max-width:920px;margin:5vh auto;max-height:88vh;display:flex;flex-direction:column">
    <div style="padding:15px 20px;border-bottom:1px solid var(--line);display:flex;align-items:center;gap:10px">
      <b style="font-size:15px" id="pvtitle">🎴 카드 미리보기</b><span class="spacer"></span>
      <span id="pvbrandsel"></span>
      <button class="btn" onclick="el('previewmodal').style.display='none'">닫기</button></div>
    <div id="pvgrid" style="padding:18px 20px;overflow:auto;display:flex;flex-wrap:wrap;gap:14px;justify-content:center"></div>
  </div></div>
<div id="acctmodal" class="modal-bg" onclick="if(event.target==this)el('acctmodal').style.display='none'">
  <div class="modal" style="max-width:560px;margin:7vh auto;max-height:86vh;display:flex;flex-direction:column">
    <div style="padding:15px 20px;border-bottom:1px solid var(--line);display:flex;align-items:center;gap:10px">
      <b style="font-size:15px">👤 Claude 계정</b>
      <span style="font-size:11.5px;color:var(--muted)">한도 차면 자동으로 다른 계정으로 전환됩니다</span><span class="spacer"></span>
      <button class="btn" onclick="el('acctmodal').style.display='none'">닫기</button></div>
    <div id="acctbody" style="padding:14px 20px;overflow:auto"></div>
    <div style="padding:12px 20px;border-top:1px solid var(--line);display:flex;gap:8px;align-items:center">
      <button class="btn pri" onclick="addAccount()">➕ 다른 Claude 계정 추가</button>
      <span style="font-size:11.5px;color:var(--muted)">검은 콘솔이 열리면 브라우저에서 그 계정으로 로그인하세요</span></div>
  </div></div>
<div id="brandmodal" class="modal-bg" onclick="if(event.target==this)closeBrandMgr()">
  <div class="modal" style="max-width:640px;margin:5vh auto;max-height:88vh;display:flex;flex-direction:column">
    <div style="padding:15px 20px;border-bottom:1px solid var(--line);display:flex;align-items:center;gap:10px">
      <b style="font-size:15px">⚙ 브랜드 관리</b><span class="spacer"></span>
      <button class="btn" onclick="closeBrandMgr()">닫기</button></div>
    <div id="brandbody" style="padding:18px 20px;overflow:auto"></div>
  </div></div>
<script>
let TABS=[], CUR=0, REC=null, seq=1, BRAND='haofactory', BRANDS=[], BFORM=null;
const el=id=>document.getElementById(id), esc=s=>(s||'').replace(/[&<>]/g,c=>({'&':'&amp;','<':'&lt;','>':'&gt;'}[c]));
function newTab(){return {id:seq++, keyword:'', folder:'', files:[], hint:'', post:null, busy:false, model:'opus'};}
// ── 브랜드 ──
function hexMix(hex,amt){hex=(hex||'#FD6F22').replace('#','');if(hex.length<6)hex='FD6F22';
  const r=parseInt(hex.substr(0,2),16),g=parseInt(hex.substr(2,2),16),b=parseInt(hex.substr(4,2),16),m=v=>Math.round(v+(255-v)*amt);
  return `rgb(${m(r)},${m(g)},${m(b)})`;}
function hexA(hex,a){hex=(hex||'#FD6F22').replace('#','');if(hex.length<6)hex='FD6F22';
  return `rgba(${parseInt(hex.substr(0,2),16)},${parseInt(hex.substr(2,2),16)},${parseInt(hex.substr(4,2),16)},${a})`;}
function curBrand(){return BRANDS.find(b=>b.id==BRAND)||{};}
function applyBrandColor(){const c=curBrand().color||'#FD6F22',R=document.documentElement;
  R.style.setProperty('--brand',c);R.style.setProperty('--brand-l',hexMix(c,0.9));R.style.setProperty('--brand-d',c);
  R.style.setProperty('--brand-2',hexMix(c,0.24));R.style.setProperty('--brand-sh',hexA(c,0.32));}
function renderBrandSel(){const s=el('brandsel');if(!s)return;
  s.innerHTML=BRANDS.map(b=>`<option value="${b.id}" ${b.id==BRAND?'selected':''}>${esc(b.name)}</option>`).join('');}
function loadBrands(cb){return fetch('/api/brands').then(r=>r.json()).then(d=>{BRANDS=d.brands||[];
  if(!BRANDS.find(b=>b.id==BRAND))BRAND=(BRANDS[0]||{}).id||'haofactory';
  renderBrandSel();applyBrandColor();if(cb)cb();});}
function switchBrand(id){BRAND=id;applyBrandColor();renderBrandSel();
  if(el('recbox')&&el('recbox').style.display=='block'){el('recbox').style.display='none';}
  toast('브랜드: '+(curBrand().name||id));}
function init(){TABS=[newTab()];CUR=0;loadBrands();render();
  fetch('/api/version').then(r=>r.json()).then(d=>{if(el('ver'))el('ver').textContent='v'+d.version;}).catch(()=>{});
  fetch('/api/status').then(r=>r.json()).then(s=>{
    if(!s.claude){el('engine').textContent='● 엔진 없음';el('engine').className='pill off';el('loginwarn').style.display='block';el('loginwarn').textContent='⚠ claude 실행파일을 찾지 못했습니다. Claude 데스크톱 앱 또는 CLI 설치 필요.';return;}
    checkAuth();
  });
}
function checkAuth(){el('engine').textContent='로그인 확인 중…';el('engine').className='pill off';
  fetch('/api/auth-status').then(r=>r.json()).then(a=>{
    if(a.logged_in){el('engine').textContent='● Claude 로그인됨';el('engine').className='pill';el('loginwarn').style.display='none';}
    else{el('engine').textContent='● 로그인 필요';el('engine').className='pill off';el('loginwarn').style.display='block';
      el('loginwarn').innerHTML='⚠ Claude 로그인이 필요합니다 (각자 자기 구독으로). <button class="btn" style="padding:5px 12px;margin-left:8px" onclick="doLogin()">🔑 로그인</button>';}
  });
}
function toast(msg,kind){const t=document.createElement('div');t.className='toast'+(kind?' '+kind:'');t.innerHTML=msg;
  el('toasts').appendChild(t);setTimeout(()=>{t.style.transition='.3s';t.style.opacity='0';t.style.transform='translateY(8px)';setTimeout(()=>t.remove(),300);}, kind=='err'?4200:2600);}
function doLogin(){fetch('/api/login',{method:'POST'}).then(r=>r.json()).then(d=>{
  if(!d.ok){alert(d.msg||'로그인 실행 실패');return;}
  el('loginwarn').innerHTML='🔑 로그인 창(검은 콘솔)이 열렸어요. 안내대로 브라우저에서 승인하면 자동으로 인식됩니다…';
  pollAuth();});}
function pollAuth(){let n=0;const iv=setInterval(()=>{n++;
  fetch('/api/auth-status').then(r=>r.json()).then(a=>{
    if(a.logged_in){clearInterval(iv);el('loginwarn').style.display='none';el('engine').textContent='✓ Claude 로그인됨';}
    else if(n>80){clearInterval(iv);el('loginwarn').innerHTML='아직 로그인 안 됨. 콘솔 창에서 완료 후 <button class="btn" style="padding:5px 12px" onclick="checkAuth()">다시 확인</button>';}
  });},3000);}
// ── 멀티 계정 ──
function openAccounts(){el('acctmodal').style.display='block';loadAccounts();}
function loadAccounts(){el('acctbody').innerHTML='<div style="color:var(--muted);font-size:13px">계정 상태 확인 중…</div>';
  fetch('/api/accounts').then(r=>r.json()).then(renderAccounts).catch(()=>{el('acctbody').innerHTML='<div class="warn">계정 정보를 불러오지 못했습니다.</div>';});}
function renderAccounts(d){
  el('acctbody').innerHTML = d.accounts.map(a=>{
    const on = a.id==d.active;
    const st = a.logged_in
      ? `<span style="color:#138a36;font-weight:700">● 로그인됨</span>${a.email?' · '+esc(a.email):''}${a.plan?' · '+esc(a.plan):''}`
      : `<span style="color:#c0392b;font-weight:700">● 로그인 필요</span>`;
    return `<div style="display:flex;align-items:center;gap:10px;padding:11px 12px;margin-bottom:8px;border:1.5px solid ${on?'var(--brand)':'var(--line)'};border-radius:11px;background:${on?'rgba(253,111,34,.05)':'#fff'}">
      <div style="flex:1;min-width:0">
        <div style="font-weight:800;font-size:13.5px">${esc(a.label)}${on?' <span style="color:var(--brand);font-size:11px">✓ 사용 중</span>':''}</div>
        <div style="font-size:11.5px;color:var(--muted);margin-top:2px">${st}</div>
      </div>
      ${!on?`<button class="btn" onclick="switchAccount('${a.id}')">전환</button>`:''}
      ${!a.logged_in?`<button class="btn pri" onclick="loginAccount('${a.id}')">🔑 로그인</button>`:''}
      ${a.id!='default'?`<button class="btn" style="color:#c0392b" onclick="removeAccount('${a.id}','${esc(a.label)}')">삭제</button>`:''}
    </div>`;}).join('');
}
function switchAccount(id){fetch('/api/account-switch',{method:'POST',headers:{'Content-Type':'application/json'},body:JSON.stringify({id})})
  .then(r=>r.json()).then(d=>{if(d.ok){toast('계정 전환됨');loadAccounts();checkAuth();}else toast(d.msg||'전환 실패','err');});}
function loginAccount(id){fetch('/api/login',{method:'POST',headers:{'Content-Type':'application/json'},body:JSON.stringify({id})})
  .then(r=>r.json()).then(d=>{if(!d.ok){toast(d.msg||'로그인 실행 실패','err');return;}
    toast('🔑 로그인 콘솔이 열렸어요');pollAccounts();});}
let PENDING_ACCT=null;
function addAccount(){const label=prompt('추가할 계정 이름 (예: 회사 계정, 두번째 구독)','추가 계정');if(label===null)return;
  fetch('/api/account-add',{method:'POST',headers:{'Content-Type':'application/json'},body:JSON.stringify({label})})
  .then(r=>r.json()).then(d=>{if(!d.ok){toast(d.msg||'추가 실패','err');return;}
    PENDING_ACCT=d.id;toast('🔑 새 계정 로그인 콘솔이 열렸어요. 로그인하면 자동 전환됩니다');loadAccounts();pollAccounts();});}
function removeAccount(id,label){if(!confirm('"'+label+'" 계정을 삭제할까요? (저장된 로그인이 지워집니다)'))return;
  fetch('/api/account-remove',{method:'POST',headers:{'Content-Type':'application/json'},body:JSON.stringify({id})})
  .then(r=>r.json()).then(d=>{if(d.ok){toast('삭제됨');loadAccounts();checkAuth();}else toast(d.msg||'삭제 실패','err');});}
function pollAccounts(){let n=0;const iv=setInterval(()=>{n++;
  fetch('/api/accounts').then(r=>r.json()).then(d=>{renderAccounts(d);
    // 로그인 기다리던 새 계정이 로그인되면 그 계정으로 자동 전환
    if(PENDING_ACCT){const p=d.accounts.find(a=>a.id==PENDING_ACCT);
      if(p&&p.logged_in){clearInterval(iv);const id=PENDING_ACCT;PENDING_ACCT=null;switchAccount(id);return;}}
    else if(d.accounts.some(a=>a.id==d.active&&a.logged_in)){clearInterval(iv);checkAuth();}
    if(n>80)clearInterval(iv);});},3000);}
function renderTabs(){
  el('tabs').innerHTML = TABS.map((t,i)=>`<div class="tab ${i==CUR?'on':''}" onclick="sel(${i})">글 ${i+1}${t.keyword?': '+esc(t.keyword.slice(0,10)):''}<span class="x" onclick="event.stopPropagation();delTab(${i})">×</span></div>`).join('')
    + `<div class="addtab" onclick="addTab()">+ 새 글</div>`;
}
function sel(i){CUR=i;render();}
function addTab(){TABS.push(newTab());CUR=TABS.length-1;render();}
function delTab(i){if(TABS.length==1)return;TABS.splice(i,1);if(CUR>=TABS.length)CUR=TABS.length-1;render();}
function t(){return TABS[CUR];}
function render(){renderTabs();
  const x=t();
  el('card').innerHTML=`
    <div class="field"><label>메인키워드</label>
      <div class="row"><input type="text" id="kw" value="${esc(x.keyword)}" placeholder="예: FRP조형물 — 입력 후 Enter로 바로 생성" oninput="t().keyword=this.value;renderTabs()" onkeydown="if(event.key=='Enter'){event.preventDefault();generate();}">
        <button class="btn" onclick="toggleRec()">📝 추천</button></div>
      <div class="recbox" id="recbox"></div>
    </div>
    <div class="field"><label>사진 폴더</label>
      <div class="row"><input type="text" id="folder" value="${esc(x.folder)}" placeholder="폴더 경로" oninput="t().folder=this.value">
        <button class="btn" onclick="pickFolder()">폴더 선택</button>
        <button class="btn" onclick="loadPhotos()">불러오기</button>
        ${x.folder&&x.files.length?`<button class="btn" onclick="organize()" title="파일명을 단계별(공정/완성 등)로 자동 정리">📷 사진 정리</button>`:''}</div>
      <div id="orgbar" style="margin-top:6px;font-size:12px;color:var(--brand);font-weight:700"></div>
      ${x.folder?`<div class="hint ${x.files.length?'ok':'bad'}">${x.files.length?('📷 사진 '+x.files.length+'장 (하위 폴더 포함) · 생성 시 키워드에 맞춰 자동 선택'):'이 폴더에 사진이 없습니다'}</div>
      <div class="photos">${x.files.slice(0,6).map(f=>`<img src="/img?folder=${encodeURIComponent(x.folder)}&name=${encodeURIComponent(f)}" title="${esc(f)}">`).join('')}${x.files.length>6?`<div style="width:62px;height:62px;border-radius:7px;background:#eef1f4;display:flex;align-items:center;justify-content:center;font-size:12px;color:#889;font-weight:700">+${x.files.length-6}</div>`:''}</div>`:''}
    </div>
    <details class="opt" ${x.hint?'open':''}><summary>＋ 프로젝트 정보 (선택 — 실제 작업 건이면 현장·소재·특이사항)</summary>
      <textarea id="hint" oninput="t().hint=this.value" placeholder="예: 전남 광양, 황소 캐릭터, 벤치 포토존, FRP">${esc(x.hint)}</textarea></details>
    ${(curBrand().card_templates||[]).length>1?`
    <div class="field"><label>카드 디자인 (클릭해서 선택)</label>
      <div class="tplrow">${(curBrand().card_templates).map(tp=>`<div class="tplopt ${(x.cardTpl||'1')==tp?'sel':''}" onclick="t().cardTpl='${tp}';render()" title="디자인 ${tp}">
        <img src="/api/template-thumb?brand=${BRAND}&template=${tp}" loading="lazy" alt="디자인 ${tp}"><span>디자인 ${tp}</span></div>`).join('')}</div></div>`:''}
    <div class="gen-row"><button class="btn pri lg" id="genbtn" onclick="generate()" ${x.busy?'disabled':''}>${x.busy?'<span class=spin></span> 생성 중…':'✍ 생성하기'}</button>
      <select onchange="t().model=this.value">
        <option value="opus" ${(x.model||'opus')=='opus'?'selected':''}>Opus (품질)</option>
        <option value="sonnet" ${x.model=='sonnet'?'selected':''}>Sonnet (빠름)</option></select>
      <span class="spacer"></span>
      <span style="font-size:12px;color:var(--ink2);font-weight:600">Opus 원고 1~2분 · 카드는 뒤이어 채워집니다</span></div>
    <div class="preview" id="preview" style="display:block">${x.post?renderPost(x.post):(x.busy?loadingState():emptyState())}</div>`;
}
function emptyState(){return `<div style="text-align:center;padding:42px 20px;color:var(--muted)">
  <div style="font-size:38px;margin-bottom:10px;opacity:.45">📝</div>
  <div style="font-weight:800;color:var(--ink2);margin-bottom:5px">아직 생성된 글이 없어요</div>
  <div style="font-size:12.5px">키워드와 사진 폴더를 정하고 <b style="color:var(--brand)">✍ 생성하기</b>를 눌러주세요</div></div>`;}
function loadingState(){return `<div style="text-align:center;padding:42px 20px;color:var(--brand)">
  <div class="spin" style="width:28px;height:28px;border-width:3px;margin-bottom:14px"></div>
  <div style="font-weight:800;color:var(--ink2)">원고를 쓰는 중…</div>
  <div style="font-size:12.5px;color:var(--muted);margin-top:4px">Opus 기준 1~2분 정도 걸려요</div></div>`;}
function toggleRec(){const b=el('recbox');if(b.style.display=='block'){b.style.display='none';return;}
  b.style.display='block';b.innerHTML='<div style="padding:14px;color:#889;font-size:12px">불러오는 중…</div>';
  fetch('/api/recommend?brand='+encodeURIComponent(BRAND)).then(r=>r.json()).then(d=>{REC=d.items||[];
    if(!REC.length){b.innerHTML='<div style="padding:16px;color:#c0392b;font-size:12.5px;line-height:1.6">'+esc(d.msg||'추천할 키워드가 없습니다.')+'</div>';return;}
    b.innerHTML='<table>'+REC.map(it=>`<tr onclick="pickKw('${esc(it.keyword).replace(/'/g,"\\'")}')"><td><b>${esc(it.keyword)}</b></td><td style="color:#889">${esc(it.category)}</td><td style="text-align:right;color:#889">${Number(it.volume).toLocaleString()}</td><td><span class="recst ${it.state}">${it.state}</span></td></tr>`).join('')+'</table>';
  }).catch(e=>{b.innerHTML='<div style="padding:16px;color:#c0392b;font-size:12.5px">추천 불러오기 오류: '+esc(''+e)+'</div>';});
}
function pickKw(k){t().keyword=k;el('recbox').style.display='none';render();}
let BR={path:'',drive:true}, BRF=[], PLACES=[];
function pickFolder(){el('brmodal').style.display='block';browse('');}
function closeBrowse(){el('brmodal').style.display='none';}
function renderPlaces(){el('brplaces').innerHTML=PLACES.map((p,i)=>`<button class="btn" style="padding:6px 11px;font-size:12px" onclick="browse(PLACES[${i}].path)">${esc(p.name)}</button>`).join('')+`<button class="btn" style="padding:6px 11px;font-size:12px" onclick="browse('')">💽 드라이브</button>`;}
function browse(path){
  fetch('/api/browse?path='+encodeURIComponent(path)).then(r=>r.json()).then(d=>{
    if(d.error){alert('열 수 없음: '+d.error);return;}
    if(d.places){PLACES=d.places;}
    renderPlaces();
    BR=d; BRF=[];
    el('brpath').textContent=d.drive?'내 컴퓨터 (드라이브 선택)':d.path;
    el('brimg').textContent=d.drive?'':('✓ 선택 시 하위 폴더까지 사진을 찾습니다'+(d.images?(' · 여기 직접 '+d.images+'장'):''));
    el('brpick').disabled=!!d.drive;
    if(!d.drive && d.parent) BRF.push({label:'↑ 상위 폴더', full:d.parent});
    d.folders.forEach(f=>{const full=d.drive?f:(d.path.replace(/[\\/]+$/,'')+'\\'+f);BRF.push({label:'📁 '+f, full:full});});
    el('brlist').innerHTML=BRF.map((it,i)=>`<div class="br-item" onclick="browse(BRF[${i}].full)">${esc(it.label)}</div>`).join('')||'<div style="padding:14px;color:#889;font-size:12px">하위 폴더 없음</div>';
  });
}
function selectFolder(){const x=t();x.folder=BR.path;closeBrowse();
  fetch('/api/photos?folder='+encodeURIComponent(BR.path)).then(r=>r.json()).then(d=>{x.files=d.files;render();});}
function loadPhotos(){const f=el('folder').value;t().folder=f;
  fetch('/api/photos?folder='+encodeURIComponent(f)).then(r=>r.json()).then(d=>{t().files=d.files;render();});}
function organize(){const x=t();
  if(!confirm('이 폴더(하위 포함) 사진을 [프로젝트]_[타입]_[단계]_[번호] 로 정리합니다.\\n· 타입=캐릭터/마스코트/글자/박스/FRP 등 (폴더명에 없으면 Claude 비전으로 폴더별 1회 판단)\\n· 단계=대표/공정/도장/설치/완성 (이름이 무의미하면 비전으로 판단)\\n타입이 파일명에 박혀 키워드 매칭이 정확해집니다.\\n원본 이름이 바뀝니다(구독 사용·시간 걸림). 진행할까요?'))return;
  fetch('/api/organize',{method:'POST',headers:{'Content-Type':'application/json'},body:JSON.stringify({folder:x.folder,brand:BRAND})})
   .then(r=>r.json()).then(d=>{if(!d.ok){alert(d.msg);return;}pollOrg();});}
function pollOrg(){const iv=setInterval(()=>{fetch('/api/organize-status').then(r=>r.json()).then(s=>{
  const last=(s.log&&s.log.length)?s.log[s.log.length-1]:'';
  if(el('orgbar'))el('orgbar').textContent='📷 정리 중… '+s.done+'/'+s.total+' · '+s.renamed+'장 변경 '+(last?('· '+last):'');
  if(!s.running&&s.finished){clearInterval(iv);if(el('orgbar'))el('orgbar').textContent='✓ 정리 완료 · '+s.renamed+'장 이름변경';loadPhotos();}
});},1500);}
function generate(){const x=t();if(!x.keyword){toast('키워드를 입력하세요','err');return;}
  x.busy=true;render();
  fetch('/api/generate',{method:'POST',headers:{'Content-Type':'application/json'},body:JSON.stringify({keyword:x.keyword,folder:x.folder,hint:x.hint,model:x.model||'opus',brand:BRAND,template:x.cardTpl||'1'})})
   .then(r=>r.json()).then(d=>{x.busy=false;
     if(!d.ok){toast(d.msg||'생성 실패','err');render();return;}
     x.post=d.post;render();toast('✍ 원고 완성','ok');
     if(d.post.cardnews_job) pollCards(d.post.cardnews_job, x);
   }).catch(e=>{x.busy=false;toast('오류: '+e,'err');render();});
}
function pollCards(job, tab){const iv=setInterval(()=>{fetch('/api/cardnews-status?id='+job).then(r=>r.json()).then(s=>{
  if(s.status=='done'){clearInterval(iv);tab.post.cardnews_pngs=s.pngs||[];tab.post.cardnews_dir=s.dir;tab.post.cardnews=s.pptx;tab.post.card_srcs=s.srcs||[];tab.post.card_state=s.cards||[];if(t()===tab)render();toast('🎴 카드뉴스 '+(s.pngs||[]).length+'장 완성','ok');}
  else if(s.status=='error'){clearInterval(iv);tab.post.cardnews_png_err=s.msg||'카드 렌더 실패';if(t()===tab)render();toast('카드 렌더 실패','err');}
}).catch(()=>{});},2000);}
function renderPost(p){
  const folder=p.folder||'', cdir=p.cardnews_dir||'';
  const pngs=p.cardnews_pngs||[], useCards=pngs.length>0;
  let pi=0;
  let body=p.blocks.map(b=>{
    if(b.type=='photo'){
      if(useCards){
        if(pi<pngs.length){ const ci=pi; const nm=pngs[ci].split(/[\\\\/]/).pop(); pi++;
          return `<div class="cardwrap"><img id="cardimg${ci}" src="/img?folder=${encodeURIComponent(cdir)}&name=${encodeURIComponent(nm)}&v=${Date.now()}">${cardCtl(ci)}</div>`; }
        pi++; return '';        // 카드 다 씀 — 여분 자리 숨김(빈칸 X)
      }
      pi++;
      if(b.file) return `<img src="/img?folder=${encodeURIComponent(folder)}&name=${encodeURIComponent(b.file)}">`;
      return `<div class="pv-photo-missing">📷 사진 자리</div>`;
    }
    return b.text.split('\n\n').map(para=>{
      const lines=para.split('\n').filter(x=>x.trim());
      if(lines.length==1 && /^\*\*.+\*\*$/.test(lines[0].trim())) return '<h3>'+esc(lines[0].replace(/\*\*/g,''))+'</h3>';
      return '<p>'+esc(para.replace(/\*\*/g,''))+'</p>';
    }).join('');
  }).join('');
  if(useCards){ while(pi<pngs.length){ const ci=pi; const nm=pngs[ci].split(/[\\\\/]/).pop(); pi++;
    body+=`<div class="cardwrap"><img id="cardimg${ci}" src="/img?folder=${encodeURIComponent(cdir)}&name=${encodeURIComponent(nm)}&v=${Date.now()}">${cardCtl(ci)}</div>`; } }
  let extra='';
  if(!pngs.length && p.cardnews_err) extra=`<div style="font-size:11px;color:#d55;margin:8px 0">카드뉴스 오류: ${esc(p.cardnews_err)}</div>`;
  else if(!pngs.length && p.cardnews_png_err) extra=`<div style="font-size:11px;color:#d55;margin:8px 0">카드 렌더 실패: ${esc(p.cardnews_png_err)}</div>`;
  else if(!pngs.length && p.cardnews_job) extra=`<div style="font-size:12px;color:var(--brand);font-weight:700;margin:8px 0"><span class="spin" style="border-color:var(--brand);border-top-color:transparent"></span> 카드뉴스 만드는 중… (원고 먼저 확인하세요. 잠시 후 본문 이미지로 채워집니다)</div>`;
  const cc=p.char_count||0, okLen=cc>=1500;
  return `<div class="pv-head"><div class="pv-title">${esc(p.title)}</div></div>
    <div class="badges">
      <span class="badge ${okLen?'ok':'warn'}">${okLen?'✓':'⚠'} 공백제외 ${cc}자${okLen?'':' · 1500 미달'}</span>
      ${useCards?`<span class="badge brand">🎴 카드뉴스 ${pngs.length}장</span>`:`<span class="badge">📷 본문사진 ${p.blocks.filter(b=>b.type=='photo'&&b.file).length}장</span>`}</div>
    <div class="naver-box">
      <div class="naver-title">📤 네이버 블로그로 옮기기</div>
      <div class="row" style="flex-wrap:wrap;gap:8px">
        <button class="btn pri" onclick="copyTitle()">① 제목 복사</button>
        <button class="btn pri" onclick="copyBody()">② 본문 복사</button>
        ${useCards?`<button class="btn" onclick="openFile(t().post.cardnews_dir)">③ 카드 이미지 폴더 열기</button>`:''}</div>
      <div class="naver-steps">네이버 글쓰기 열기 → <b>①</b> 제목칸에 붙여넣기 → <b>②</b> 본문에 붙여넣기 → <b>③</b> 폴더 열고 <b>${useCards?'[사진N] 자리에 card 1~'+pngs.length+'을 순서대로':'사진을'}</b> 드래그</div>
    </div>
    <div class="row" style="margin:10px 0;flex-wrap:wrap">
      <button class="btn" onclick="saveDocx()">📄 워드로 저장</button>
      <button class="btn" onclick="generate()">↻ 다시 생성</button></div>
    ${useCards?`<div class="cardsec-hint">🖱 카드에 마우스를 올리면 하단 툴바에서 <b style="color:var(--ink2)">✏️ 글 수정 · 🔄 사진 교체 · 위치 ▲▼◀▶ · 확대 ＋－</b></div>`:''}
    ${extra}
    <div class="pv-body">${body}</div>`;
}
function copyClip(text,msg){const done=()=>toast(msg,'ok');
  if(navigator.clipboard&&navigator.clipboard.writeText){navigator.clipboard.writeText(text).then(done).catch(()=>fallbackCopy(text,done));}
  else fallbackCopy(text,done);}
function fallbackCopy(text,done){const ta=document.createElement('textarea');ta.value=text;ta.style.position='fixed';ta.style.opacity='0';document.body.appendChild(ta);ta.select();try{document.execCommand('copy');done();}catch(e){toast('복사 실패','err');}document.body.removeChild(ta);}
function copyTitle(){copyClip(t().post.title||'','① 제목 복사됨 · 네이버 제목칸에 붙여넣기');}
function copyBody(){const p=t().post;const pngs=p.cardnews_pngs||[],useCards=pngs.length>0;let pi=0;const out=[];
  p.blocks.forEach(b=>{
    if(b.type=='photo'){
      if(useCards){if(pi<pngs.length){pi++;out.push('📷 [사진'+pi+']');}else pi++;}
      else out.push('📷 [사진]');
    } else { b.text.split('\n\n').forEach(para=>{const tx=para.replace(/\*\*/g,'').trim();if(tx)out.push(tx);}); }
  });
  if(useCards){while(pi<pngs.length){pi++;out.push('📷 [사진'+pi+']');}}
  copyClip(out.join('\n\n'),'② 본문 복사됨 · 네이버 본문에 붙여넣기 ([사진N] 자리에 카드 드래그)');}
function saveDocx(){toast('📄 워드 저장 중…');fetch('/api/docx',{method:'POST',headers:{'Content-Type':'application/json'},body:JSON.stringify({post:t().post})})
  .then(r=>r.json()).then(d=>{if(d.ok)toast('📄 워드 저장됨 · 파일이 열립니다','ok');else toast(d.msg||'저장 실패','err');});}
function openFile(p){fetch('/api/open',{method:'POST',headers:{'Content-Type':'application/json'},body:JSON.stringify({path:p})}).then(r=>r.json()).then(d=>{if(!d.ok)toast('파일을 열 수 없습니다','err');});}
let SWAP_CI=0, SWAP_FILES=[];
function openSwap(ci){SWAP_CI=ci;el('swapmodal').style.display='block';
  el('swaplist').innerHTML='<div style="padding:14px;color:#889;font-size:12px">사진 불러오는 중…</div>';
  fetch('/api/photos?folder='+encodeURIComponent(t().folder)).then(r=>r.json()).then(d=>{SWAP_FILES=d.files;
    el('swaplist').innerHTML=d.files.map((f,i)=>`<img src="/img?folder=${encodeURIComponent(t().folder)}&name=${encodeURIComponent(f)}" title="${esc(f)}" onclick="doSwap(${i})" style="width:84px;height:84px;object-fit:cover;border-radius:8px;cursor:pointer;border:2px solid transparent">`).join('')||'<div style="padding:14px;color:#889">사진 없음</div>';});}
function closeSwap(){el('swapmodal').style.display='none';}
function doSwap(i){const file=SWAP_FILES[i], x=t().post;
  el('swaplist').style.opacity='.5';
  fetch('/api/swap-card',{method:'POST',headers:{'Content-Type':'application/json'},body:JSON.stringify({pptx:x.cardnews,cardnews_dir:x.cardnews_dir,folder:x.folder,index:SWAP_CI,file:file})})
   .then(r=>r.json()).then(d=>{el('swaplist').style.opacity='1';
     if(!d.ok){toast(d.msg||'교체 실패','err');return;}
     closeSwap();
     (x.card_srcs=x.card_srcs||[])[SWAP_CI]=d.src||'';      // 교체된 원본 기억(위치조정용)
     (x.card_adj=x.card_adj||{})[SWAP_CI]={cx:.5,cy:.5,zoom:1};   // 위치/확대 초기화
     const img=el('cardimg'+SWAP_CI); if(img) img.src='/img?folder='+encodeURIComponent(x.cardnews_dir)+'&name='+encodeURIComponent(d.name)+'&v='+Date.now();
     toast('🔄 사진 교체됨','ok');
   }).catch(e=>{el('swaplist').style.opacity='1';toast('오류: '+e,'err');});}
// ── 카드별 위치 이동 / 확대 ──
function cardCtl(ci){return `<div class="cardbar">
  <button class="tb" onclick="openTextEdit(${ci})" title="카드 글 수정">✏️ 글</button>
  <button class="tb" onclick="openSwap(${ci})" title="사진 교체">🔄 사진</button>
  <span class="sep"></span>
  <button onclick="nudge(${ci},0,-.08)" title="위로">▲</button>
  <button onclick="nudge(${ci},-.08,0)" title="왼쪽">◀</button>
  <button onclick="resetCard(${ci})" title="가운데/초기화">⟳</button>
  <button onclick="nudge(${ci},.08,0)" title="오른쪽">▶</button>
  <button onclick="nudge(${ci},0,.08)" title="아래로">▼</button>
  <span class="sep"></span>
  <button onclick="zoomCard(${ci},.15)" title="확대">＋</button>
  <button onclick="zoomCard(${ci},-.15)" title="축소">－</button></div>`;}
function _adj(ci){const x=t().post;const A=(x.card_adj=x.card_adj||{});return A[ci]||(A[ci]={cx:.5,cy:.5,zoom:1});}
function nudge(ci,dx,dy){const a=_adj(ci);a.cx=Math.max(0,Math.min(1,a.cx+dx));a.cy=Math.max(0,Math.min(1,a.cy+dy));adjustCard(ci);}
function zoomCard(ci,dz){const a=_adj(ci);a.zoom=Math.max(1,Math.min(2.5,(a.zoom||1)+dz));adjustCard(ci);}
function resetCard(ci){const x=t().post;(x.card_adj=x.card_adj||{})[ci]={cx:.5,cy:.5,zoom:1};adjustCard(ci);}
function adjustCard(ci){const x=t().post;const a=_adj(ci);const src=(x.card_srcs||[])[ci]||'';
  if(!src){toast('이 카드는 [🔄 사진 교체]를 한 번 누른 뒤 조정돼요','err');return;}
  const img=el('cardimg'+ci); if(img)img.style.opacity='.4';
  fetch('/api/adjust-card',{method:'POST',headers:{'Content-Type':'application/json'},body:JSON.stringify({pptx:x.cardnews,cardnews_dir:x.cardnews_dir,src:src,index:ci,cx:a.cx,cy:a.cy,zoom:a.zoom})})
   .then(r=>r.json()).then(d=>{if(img)img.style.opacity='1';
     if(!d.ok){toast(d.msg||'조정 실패','err');return;}
     if(img)img.src='/img?folder='+encodeURIComponent(x.cardnews_dir)+'&name='+encodeURIComponent(d.name)+'&v='+Date.now();
   }).catch(e=>{if(img)img.style.opacity='1';toast('오류: '+e,'err');});}
// ── 카드 글 수정 ──
let TXT_CI=0;
function openTextEdit(ci){TXT_CI=ci;const x=t().post;const c=(x.card_state||[])[ci]||{};
  const flds=c.fields||[];const lbl={title:'제목 (블로그 글 제목)',headline:'헤드라인',subtitle:'부제',body:'본문 설명'};
  let html='';
  flds.forEach(f=>{const big=(f=='body'||f=='title');
    html+=`<div class="field"><label>${lbl[f]||f}</label>${big?`<textarea id="tx_${f}" style="height:64px">${esc(c[f]||'')}</textarea>`:`<input type="text" id="tx_${f}" value="${esc(c[f]||'')}">`}</div>`;});
  if(!flds.length) html='<div style="color:#889;font-size:13px;padding:6px 0">이 카드는 글자가 없는 디자인이에요 (사진만). 🔄 사진 교체를 써주세요.</div>';
  el('txtfields').innerHTML=html;el('textmodal').style.display='block';
  const first=el('tx_'+(flds[0]||'')); if(first)setTimeout(()=>first.focus(),50);}
function applyText(){const x=t().post;const c=(x.card_state||[])[TXT_CI]||{};const flds=c.fields||[];
  const body={cardnews_dir:x.cardnews_dir,index:TXT_CI};
  flds.forEach(f=>{const e=el('tx_'+f);if(e){body[f]=e.value;c[f]=e.value;}});
  fetch('/api/card-text',{method:'POST',headers:{'Content-Type':'application/json'},body:JSON.stringify(body)})
   .then(r=>r.json()).then(d=>{if(!d.ok){toast(d.msg||'수정 실패','err');return;}
     const img=el('cardimg'+TXT_CI);if(img)img.src='/img?folder='+encodeURIComponent(x.cardnews_dir)+'&name='+encodeURIComponent(d.name)+'&v='+Date.now();
     el('textmodal').style.display='none';toast('✏️ 글 수정됨','ok');}).catch(e=>toast('오류: '+e,'err'));}
// ── 카드 미리보기 (최종 체크용) ──
function pvBrandOptions(cur){return '<select onchange="openCardPreview(this.value)" style="border:1.5px solid var(--line);border-radius:9px;padding:6px 9px;font-size:12px;font-weight:700">'+
  BRANDS.filter(b=>b.has_cards).map(b=>`<option value="${b.id}" ${b.id==cur?'selected':''}>${esc(b.name)}</option>`).join('')+'</select>';}
function previewTpl(){openCardPreview(BRAND,(t().cardTpl||'1'));}
function openCardPreview(bid,tpl){tpl=tpl||'1';el('previewmodal').style.display='block';el('pvbrandsel').innerHTML=pvBrandOptions(bid);
  el('pvtitle').textContent='🎴 카드 미리보기';
  el('pvgrid').innerHTML='<div style="padding:44px;color:#889;font-size:13px"><span class="spin" style="border-color:var(--brand);border-top-color:transparent;margin-right:8px"></span>샘플 사진으로 카드 그리는 중…</div>';
  const folder=(t()&&t().folder)||'';
  fetch('/api/preview-cards',{method:'POST',headers:{'Content-Type':'application/json'},body:JSON.stringify({brand:bid,folder:folder,template:tpl})})
   .then(r=>r.json()).then(d=>{if(!d.ok){el('pvgrid').innerHTML='<div style="padding:24px;color:#d55">'+esc(d.msg||'실패')+'</div>';return;}
     el('pvtitle').textContent='🎴 '+d.name+' 카드 미리보기 ('+d.pngs.length+'장)';
     el('pvgrid').innerHTML=d.pngs.map((p,i)=>{const nm=p.split(/[\\\\/]/).pop();
       return '<div style="text-align:center"><img src="/img?folder='+encodeURIComponent(d.dir)+'&name='+encodeURIComponent(nm)+'&v='+Date.now()+'" style="width:224px;height:224px;object-fit:cover;border-radius:13px;box-shadow:var(--sh)"><div style="font-size:11px;color:#889;margin-top:5px;font-weight:700">'+(i==0?'표지':'카드 '+(i+1))+'</div></div>';}).join('');
   }).catch(e=>{el('pvgrid').innerHTML='<div style="padding:24px;color:#d55">오류: '+e+'</div>';});}
// ── 브랜드 관리 모달 ──
function openBrandMgr(){el('brandmodal').style.display='block';BFORM=null;loadBrands(renderBrandMgr);}
function closeBrandMgr(){el('brandmodal').style.display='none';loadBrands();}
function renderBrandMgr(){
  if(BFORM){renderBrandForm();return;}
  el('brandbody').innerHTML=`<div style="display:flex;gap:8px;flex-wrap:wrap;margin-bottom:12px">
    ${BRANDS.map(b=>`<button class="btn" onclick="editBrand('${b.id}')"><span style="display:inline-block;width:10px;height:10px;border-radius:50%;background:${b.color};margin-right:7px;vertical-align:-1px"></span>${esc(b.name)} ${b.has_cards?'🎴':''}</button>`).join('')}
    <button class="btn pri" onclick="newBrand()">+ 새 브랜드</button></div>
    <div class="row" style="margin-bottom:12px;flex-wrap:wrap"><button class="btn" onclick="syncBrands()">🔗 통검 브랜드 동기화</button>
      <button class="btn pri" onclick="openCardPreview((BRANDS.find(b=>b.has_cards)||{}).id||'haofactory')">🎴 카드 미리보기 (최종 체크)</button></div>
    <div style="font-size:12px;color:var(--muted)">브랜드를 클릭하면 설정을 수정합니다 · 🎴 = 카드 템플릿 등록됨</div>`;
}
function syncBrands(){toast('🔗 통검 브랜드 동기화 중…');
  fetch('/api/sync-brands',{method:'POST'}).then(r=>r.json()).then(d=>{
    toast('🔗 동기화 완료 · 추가 '+(d.added||[]).length+'개 · 연결 '+(d.linked||[]).length+'개','ok');
    loadBrands(renderBrandMgr);}).catch(e=>toast('오류: '+e,'err'));}
function newBrand(){BFORM={id:'',name:'',homepage:'',color:'#3B82F6',industry:'',tone:'',sections:'',extra:'',cta:'',type_words:'',card_headlines:'',tonggeom:''};renderBrandForm();}
function editBrand(id){fetch('/api/brand?id='+id).then(r=>r.json()).then(b=>{
  BFORM=Object.assign({},b,{type_words:(b.type_words||[]).join(', '),card_headlines:(b.card_headlines||[]).join('\n')});renderBrandForm();});}
function renderBrandForm(){const b=BFORM;
  const fi=(label,key,ph)=>`<div class="field"><label>${label}</label><input type="text" id="bf_${key}" value="${esc(b[key]||'')}" placeholder="${esc(ph||'')}"></div>`;
  const fa=(label,key,ph)=>`<div class="field"><label>${label}</label><textarea id="bf_${key}" placeholder="${esc(ph||'')}">${esc(b[key]||'')}</textarea></div>`;
  el('brandbody').innerHTML=`
    <div class="row" style="margin-bottom:12px"><button class="btn" onclick="BFORM=null;renderBrandMgr()">← 목록</button>
      <b style="font-size:14px">${b.id?esc(b.name)+' 수정':'새 브랜드'}</b></div>
    <div class="row" style="align-items:flex-end">${fi('브랜드 이름 *','name','예: 퍼스트디자인')}
      <div class="field" style="width:130px"><label>브랜드색</label><input type="text" id="bf_color" value="${esc(b.color||'#3B82F6')}" placeholder="#3B82F6"></div></div>
    <div class="row">${fi('글 속 표기명 (비우면 이름)','label','예: FIRST DESIGN')}
      ${fi('홈페이지 (추천키워드 기준)','homepage','firstd.co.kr')}</div>
    ${fi('업종 (한 줄)','industry','예: 기업 디자인 파트너')}
    <div style="font-size:11.5px;color:var(--muted);font-weight:700;margin:6px 0 2px">▼ 구조(소제목6·제목키워드·SEO)는 전 브랜드 공통. 아래 톤·내용만 브랜드별로 다릅니다.</div>
    ${fa('🏢 회사 소개 · 서비스 · 강점','identity','예: OO는 기획부터 제작까지 원스톱으로 진행하는 회사다. 카탈로그·포스터·패키지·촬영을 다루고 다국어 디자인이 가능하다. ...')}
    ${fa('🗣 문체/톤','tone','예: 친근한 구어체 ~해요/:). 고객 말 큰따옴표 인용. (또는) 정중·신뢰 톤, 이모티콘 절제.')}
    ${fa('🏁 마무리 CTA','cta','예: 메인키워드 제대로 하고 싶다면 OO에서 시작해보세요! 식으로 마무리.')}
    <div style="border-top:1px solid var(--line2);margin:10px 0 12px"></div>
    ${fi('타입 단어 (콤마 — 사진 매칭·정리용)','type_words','예: 카탈로그, 포스터, 전단지, 명함, 패키지')}
    ${fa('카드 헤드라인 7줄 (선택)','card_headlines','표지\\n후킹\\n이유\\n강점\\n품질\\n과정\\n마무리')}
    ${fi('통검체크 dist 폴더 (선택 — 동기화로 자동입력)','tonggeom','C:/Users/.../브랜드 통검체크')}
    <details style="margin-top:10px"><summary style="font-size:12px;color:var(--muted);cursor:pointer">고급: 완전 커스텀 프롬프트 (있으면 위 설정 대신 이걸 사용)</summary>
      <textarea id="bf_prompt" style="height:120px;font-size:12px;margin-top:8px" placeholder="구조까지 완전히 직접 쓰고 싶을 때만. 비우면 공통 구조 + 위 톤/회사소개를 사용합니다.">${esc(b.prompt||'')}</textarea></details>
    <div class="row" style="margin-top:10px;flex-wrap:wrap"><button class="btn pri" onclick="saveBrand()">💾 저장</button>
      ${b.id&&b.has_cards?`<button class="btn" onclick="openCardPreview('${b.id}')">👁 카드 미리보기</button>`:''}</div>
    ${b.id?`<div style="margin-top:10px;display:flex;gap:6px;flex-wrap:wrap;align-items:center">
      <span style="font-size:12px;color:var(--muted);font-weight:700">카드 디자인 (PPTX, 여러 개면 생성 시 선택):</span>
      ${['1','2','3'].map(s=>{const has=(b.card_templates||[]).includes(s);return `<button class="btn" onclick="openTemplate('${b.id}','${s}')">${has?'🎴':'＋'} 디자인 ${s} ${has?'교체':'등록'}</button>`;}).join('')}
     </div>`:'<div style="margin-top:10px;font-size:12px;color:var(--muted)">먼저 저장하면 카드 템플릿(PPTX)을 등록할 수 있어요</div>'}`;
}
function bval(k){const e=el('bf_'+k);return e?e.value:'';}
function saveBrand(){const cfg={id:BFORM.id||'',name:bval('name'),label:bval('label'),color:bval('color'),homepage:bval('homepage'),industry:bval('industry'),
   identity:bval('identity'),tone:bval('tone'),cta:bval('cta'),prompt:bval('prompt'),type_words:bval('type_words'),card_headlines:bval('card_headlines'),tonggeom:bval('tonggeom'),stages:BFORM.stages};
  if(!cfg.name.trim()){toast('브랜드 이름을 입력하세요','err');return;}
  fetch('/api/brand-save',{method:'POST',headers:{'Content-Type':'application/json'},body:JSON.stringify(cfg)})
   .then(r=>r.json()).then(d=>{if(!d.ok){toast(d.msg||'저장 실패','err');return;}
     toast('💾 저장됨: '+d.brand.name,'ok');
     BFORM=Object.assign({},d.brand,{type_words:(d.brand.type_words||[]).join(', '),card_headlines:(d.brand.card_headlines||[]).join('\n')});
     BRAND=d.id;loadBrands(renderBrandForm);});}
function openTemplate(id,slot){slot=slot||'1';const inp=document.createElement('input');inp.type='file';inp.accept='.pptx';
  inp.onchange=()=>{const f=inp.files[0];if(!f)return;toast('🎴 디자인 '+slot+' 업로드·추출 중…');
    const fd=new FormData();fd.append('id',id);fd.append('slot',slot);fd.append('file',f);
    fetch('/api/brand-template-file',{method:'POST',body:fd}).then(r=>r.json()).then(d=>{
      if(!d.ok){toast(d.msg||'추출 실패','err');return;}
      toast('🎴 디자인 '+(d.slot||slot)+' 등록됨 ('+d.slides+'장, 장식 '+d.assets+'개)','ok');loadBrands(()=>editBrand(id));});};
  inp.click();}
init();
</script></body></html>"""

URL = "http://127.0.0.1:5002"


def main():
    """런처(launcher.py) 진입점 — 자동업데이트 후 이 함수를 호출."""
    if "--noopen" not in sys.argv:
        threading.Thread(target=lambda: (time.sleep(1.2), webbrowser.open(URL)), daemon=True).start()
    print("블로그 작성기 →", URL, " (engine:", CLAUDE or "없음", ")")
    app.run(host="127.0.0.1", port=5002, threaded=True)


if __name__ == "__main__":
    main()
